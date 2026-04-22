from io import BytesIO
import json
import re

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph
from PyPDF2 import PdfReader
import requests


DEFAULT_MODEL = "phi3:mini"
OLLAMA_CHAT_URL = "http://localhost:11434/api/chat"
OLLAMA_TAGS_URL = "http://localhost:11434/api/tags"

KNOWN_SECTION_HEADERS = {
    "summary",
    "profile",
    "objective",
    "skills",
    "technical skills",
    "core skills",
    "projects",
    "project",
    "experience",
    "work experience",
    "internship",
    "internships",
    "education",
    "certifications",
    "achievements",
    "publications",
    "positions of responsibility",
    "leadership",
    "extracurricular",
    "languages",
    "interests",
}


def extract_resume_text(uploaded_file):
    name = uploaded_file.name.lower()
    raw = uploaded_file.getvalue()

    if name.endswith(".txt"):
        try:
            return raw.decode("utf-8")
        except UnicodeDecodeError:
            return raw.decode("latin-1")

    if name.endswith(".docx"):
        document = docx.Document(BytesIO(raw))
        chunks = []
        for para in document.paragraphs:
            if para.text.strip():
                chunks.append(para.text.strip())
        return "\n".join(chunks).strip()

    if name.endswith(".pdf"):
        reader = PdfReader(BytesIO(raw))
        chunks = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                chunks.append(page_text.strip())
        return "\n".join(chunks).strip()

    raise ValueError("Unsupported resume format. Use DOCX, PDF, or TXT.")


def check_ollama_health(timeout_sec=5):
    try:
        resp = requests.get(OLLAMA_TAGS_URL, timeout=timeout_sec)
        resp.raise_for_status()
        data = resp.json()
        models = [m.get("name", "") for m in data.get("models", [])]
        return True, models
    except Exception:
        return False, []


def warmup_model(model=DEFAULT_MODEL, timeout_sec=300):
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": "Reply only: ready"}],
        "stream": False,
        "keep_alive": "30m",
        "options": {"temperature": 0, "num_predict": 8, "num_ctx": 1024},
    }
    try:
        response = requests.post(OLLAMA_CHAT_URL, json=payload, timeout=(10, timeout_sec))
        response.raise_for_status()
        data = response.json()
        return data.get("message", {}).get("content", "").strip()
    except requests.Timeout as exc:
        raise RuntimeError(
            f"Model warm-up timed out after {timeout_sec}s. "
            "Increase timeout for first load on slow CPU."
        ) from exc


def _profile_options(speed_profile):
    profile = str(speed_profile or "fast").lower()
    if profile == "detailed":
        return {
            "temperature": 0.2,
            "num_ctx": 8192,
            "num_predict": 900,
            "max_section_chars": 9000,
            "max_jd_chars": 9000,
        }
    if profile == "balanced":
        return {
            "temperature": 0.15,
            "num_ctx": 4096,
            "num_predict": 600,
            "max_section_chars": 6000,
            "max_jd_chars": 6000,
        }
    return {
        "temperature": 0.1,
        "num_ctx": 2048,
        "num_predict": 350,
        "max_section_chars": 3500,
        "max_jd_chars": 3500,
    }


def _trim_text(text, max_chars):
    raw = str(text or "").replace("\r\n", "\n")
    lines = [ln.strip() for ln in raw.split("\n")]
    cleaned = "\n".join(ln for ln in lines if ln)
    return cleaned[:max_chars]


def _set_paragraph_text_preserve_style(paragraph, new_text):
    text = str(new_text)
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
        return

    # Preserve mixed run formatting pattern by distributing new text
    # proportionally across the existing runs.
    weights = [max(len(run.text), 1) for run in runs]
    total = sum(weights) or len(runs)
    consumed = 0
    for idx, run in enumerate(runs):
        if idx == len(runs) - 1:
            run.text = text[consumed:]
            continue
        share = round(len(text) * (weights[idx] / total))
        min_remaining = len(runs) - idx - 1
        share = max(0, min(share, len(text) - consumed - min_remaining))
        run.text = text[consumed:consumed + share]
        consumed += share


def _iter_block_items(document):
    for child in document.element.body.iterchildren():
        tag = child.tag.lower()
        if tag.endswith("}p"):
            yield "paragraph", Paragraph(child, document)
        elif tag.endswith("}tbl"):
            yield "table", Table(child, document)


def _table_to_lines(table):
    lines = []
    for row in table.rows:
        cells = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells]
        lines.append(" | ".join(cells))
    return lines


def _set_cell_text_preserve_style(cell, text):
    if not cell.paragraphs:
        cell.add_paragraph(str(text))
        return
    _set_paragraph_text_preserve_style(cell.paragraphs[0], str(text))
    for para in cell.paragraphs[1:]:
        _set_paragraph_text_preserve_style(para, "")


def _slug(text):
    slug = re.sub(r"[^a-z0-9]+", "_", str(text).lower()).strip("_")
    return slug or "section"


def _looks_like_heading(paragraph, text):
    stripped = str(text).strip()
    if not stripped:
        return False

    style_name = ""
    try:
        style_name = (paragraph.style.name or "").lower()
    except Exception:
        style_name = ""

    if "heading" in style_name:
        return True

    compact = re.sub(r"[^a-zA-Z ]", "", stripped).strip().lower()
    if compact in KNOWN_SECTION_HEADERS:
        return True

    words = stripped.replace(":", "").split()
    if stripped.endswith(":") and len(words) <= 5:
        return True
    if stripped.isupper() and len(words) <= 5:
        return True
    return False


def extract_docx_sections(docx_bytes):
    document = docx.Document(BytesIO(docx_bytes))
    blocks = list(_iter_block_items(document))

    sections = []
    current_title = "Profile"
    current_blocks = []
    current_lines = []
    section_counter = 1

    def flush_current():
        nonlocal section_counter, current_blocks, current_lines, current_title
        if not current_blocks:
            return
        section_id = f"{_slug(current_title)}_{section_counter}"
        sections.append({
            "id": section_id,
            "title": current_title,
            "blocks": list(current_blocks),
            "text": "\n".join(current_lines).strip(),
        })
        section_counter += 1
        current_blocks = []
        current_lines = []

    for block_index, (kind, item) in enumerate(blocks):
        if kind == "paragraph":
            text = item.text.strip()
            if not text:
                continue
            if _looks_like_heading(item, text):
                flush_current()
                current_title = text.rstrip(":")
                continue
            current_blocks.append({"type": "paragraph", "index": block_index, "line_count": 1})
            current_lines.append(text)
        else:
            table_lines = _table_to_lines(item)
            if not table_lines:
                continue
            current_blocks.append({"type": "table", "index": block_index, "line_count": len(table_lines)})
            current_lines.extend(table_lines)

    flush_current()

    if sections:
        return sections

    fallback_blocks = []
    fallback_lines = []
    for block_index, (kind, item) in enumerate(blocks):
        if kind == "paragraph":
            text = item.text.strip()
            if text:
                fallback_blocks.append({"type": "paragraph", "index": block_index, "line_count": 1})
                fallback_lines.append(text)
        else:
            lines = _table_to_lines(item)
            if lines:
                fallback_blocks.append({"type": "table", "index": block_index, "line_count": len(lines)})
                fallback_lines.extend(lines)

    return [{
        "id": "resume_1",
        "title": "Resume",
        "blocks": fallback_blocks,
        "text": "\n".join(fallback_lines).strip(),
    }]


def _build_section_prompt(section_title, section_text, jd_text, instructions):
    return f"""
You are editing one section of a resume for ATS alignment.

Rules:
- Keep all facts truthful to the provided section content.
- Do not invent new projects, skills, dates, companies, or achievements.
- Improve clarity, action verbs, and keyword alignment naturally.
- Keep concise resume language.
- Return only the rewritten section text (plain text).

Target section: {section_title}

Job description / criteria:
{jd_text}

User instructions:
{instructions}

Original section text:
{section_text}
""".strip()


def tailor_section_with_phi(
    section_title,
    section_text,
    jd_text="",
    instructions="",
    model=DEFAULT_MODEL,
    timeout_sec=300,
    speed_profile="fast",
):
    profile = _profile_options(speed_profile)
    prompt = _build_section_prompt(
        section_title=section_title,
        section_text=_trim_text(section_text, profile["max_section_chars"]),
        jd_text=_trim_text(jd_text, profile["max_jd_chars"]),
        instructions=instructions,
    )

    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "stream": False,
        "keep_alive": "30m",
        "options": {
            "temperature": profile["temperature"],
            "num_ctx": profile["num_ctx"],
            "num_predict": profile["num_predict"],
        },
    }

    try:
        response = requests.post(OLLAMA_CHAT_URL, json=payload, timeout=(10, timeout_sec))
        response.raise_for_status()
    except requests.Timeout as exc:
        raise RuntimeError(
            f"Section generation timed out after {timeout_sec}s. "
            "Increase timeout or keep speed profile on fast."
        ) from exc
    except requests.RequestException as exc:
        raise RuntimeError(
            "Could not connect to Ollama at http://localhost:11434. "
            "Ensure `ollama serve` is running."
        ) from exc

    try:
        data = response.json()
    except json.JSONDecodeError as exc:
        raise RuntimeError("Invalid JSON response from Ollama.") from exc

    content = data.get("message", {}).get("content") if isinstance(data, dict) else ""
    content = (content or "").strip()
    if not content:
        raise RuntimeError("Ollama returned empty content for this section.")
    return content


def apply_section_texts_to_docx(original_docx_bytes, sections, edited_section_map):
    document = docx.Document(BytesIO(original_docx_bytes))
    blocks = list(_iter_block_items(document))

    for section in sections:
        section_id = section["id"]
        section_blocks = section.get("blocks", [])
        if not section_blocks:
            continue

        edited_text = str(edited_section_map.get(section_id, section.get("text", "")))
        edited_lines = [ln.strip() for ln in edited_text.splitlines() if ln.strip() and not ln.strip().startswith("Rewritten section text")]
        if not edited_lines:
            edited_lines = [""]

        cursor = 0
        paragraph_block_indices = []

        for block in section_blocks:
            btype = block.get("type")
            bidx = block.get("index")
            line_count = max(int(block.get("line_count", 1)), 1)
            chunk = edited_lines[cursor:cursor + line_count]
            cursor += line_count

            if bidx >= len(blocks):
                continue
            kind, item = blocks[bidx]

            if btype == "paragraph" and kind == "paragraph":
                paragraph_block_indices.append(bidx)
                _set_paragraph_text_preserve_style(item, chunk[0] if chunk else "")
            elif btype == "table" and kind == "table":
                if not chunk:
                    continue
                for row_idx, row in enumerate(item.rows):
                    if row_idx >= len(chunk):
                        break
                    row_line = chunk[row_idx]
                    columns = [cell.strip() for cell in row_line.split("|")]
                    for col_idx, cell in enumerate(row.cells):
                        value = columns[col_idx] if col_idx < len(columns) else ""
                        _set_cell_text_preserve_style(cell, value)

        if cursor < len(edited_lines):
            overflow_lines = edited_lines[cursor:]
            overflow_text = "\n".join(overflow_lines).strip()
            if overflow_text and paragraph_block_indices:
                _, para = blocks[paragraph_block_indices[-1]]
                merged = para.text.strip()
                merged = f"{merged}\n{overflow_text}" if merged else overflow_text
                _set_paragraph_text_preserve_style(para, merged)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.read()
