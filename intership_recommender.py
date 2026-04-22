import streamlit as st
import pandas as pd
import numpy as np
from PyPDF2 import PdfReader
import re
import nltk
from nltk.stem import WordNetLemmatizer
from nltk.corpus import wordnet
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from bs4 import BeautifulSoup
from io import BytesIO
import base64
import json
import requests
import docx
from company_intelligence import get_company_info

ADJ, ADJ_SAT, ADV, NOUN, VERB = 'a', 's', 'r', 'n', 'v'
POS_LIST = [NOUN, VERB, ADJ, ADV]
NUM_POSTING = 50
TOP_N_KEYWORDS = 5

# ══════════════════════════════════════════════════════════════════════════════
# LOCAL LLM (OLLAMA)
# ══════════════════════════════════════════════════════════════════════════════

def _call_local_llm(prompt: str, model: str = "phi3:mini") -> str:
    """Call local Ollama model and return generated text."""
    try:
        response = requests.post(
            "http://localhost:11434/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "stream": False
            },
            timeout=60
        )
        response.raise_for_status()
        data = response.json()
        return data.get("response", "").strip()
    except requests.exceptions.RequestException as e:
        return (
            f'{{"technical":["Error connecting to Ollama: {str(e)}"],'
            f'"behavioral":[],"company_specific":[],"tips":[],"resources":[]}}'
        )


# ══════════════════════════════════════════════════════════════════════════════
# FILE HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def get_base64_image(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode()


def read_pdf(file):
    try:
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return ""

def read_docx(file):
    try:
        file_bytes = BytesIO(file.read())
        document = docx.Document(file_bytes)
        text = ""
        for para in document.paragraphs:
            if para.text.strip():
                text += para.text.strip() + "\n"
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text.strip() + " "
            text += "\n"
        return text.strip()
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
        return ""

def read_txt(file):
    try:
        raw = file.read()
        try:
            return raw.decode("utf-8").strip()
        except UnicodeDecodeError:
            return raw.decode("latin-1").strip()
    except Exception as e:
        st.error(f"Error reading TXT: {e}")
        return ""

def read_resume(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):    return read_pdf(uploaded_file)
    elif name.endswith(".docx"): return read_docx(uploaded_file)
    elif name.endswith(".txt"):  return read_txt(uploaded_file)
    else:
        st.error("Unsupported file type.")
        return ""


# ══════════════════════════════════════════════════════════════════════════════
# JOB TABLE WITH TOOLTIPS
# ══════════════════════════════════════════════════════════════════════════════

def render_table_with_tooltips(df: pd.DataFrame):
    rows_data = []
    for _, row in df.iterrows():
        company  = str(row.get("Company",  "")).strip()
        role     = str(row.get("Role",     "")).strip()
        location = str(row.get("Location", "")).strip()
        date_p   = str(row.get("Date Posted", "")).strip()
        raw_link = str(row.get("Application/Link", "#")).strip()

        url_match = re.search(r'href=["\']([^"\']+)["\']', raw_link)
        url = url_match.group(1) if url_match else raw_link

        if not company or company in ('↳', ''):
            company = rows_data[-1]["company"] if rows_data else ""

        intel = get_company_info(company, role)

        rows_data.append({
            "company":  company,
            "role":     role,
            "location": location,
            "date":     date_p,
            "url":      url,
            "intel": {
                "rating":       intel["rating"],
                "rating_src":   intel.get("rating_src", "AmbitionBox"),
                "salary":       intel["salary"],
                "size":         intel["size"],
                "employees":    intel.get("employees", "—"),
                "hq":           intel.get("hq", "India"),
                "founded":      intel.get("founded", "—"),
                "requirements": intel["requirements"][:6],
            }
        })

    rows_json = json.dumps(rows_data, ensure_ascii=False)

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  * {{ box-sizing:border-box; margin:0; padding:0; }}
  body {{ background:transparent; font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif; color:#e8eef8; }}

  /* ── Table ───────────────────────────────────────────── */
  #job-table {{ width:100%; border-collapse:collapse;
    background:rgba(0,0,0,0.72); border-radius:14px; overflow:hidden;
    box-shadow:0 6px 32px rgba(0,0,0,0.55); }}
  #job-table thead tr {{ background:#000; border-bottom:2px solid rgba(74,144,217,0.5); }}
  #job-table thead th {{ padding:11px 14px; color:#7ec8f7; font-weight:700;
    font-size:0.8rem; text-transform:uppercase; letter-spacing:0.5px; text-align:left; }}
  #job-table tbody tr {{ border-bottom:1px solid rgba(74,144,217,0.12); transition:background 0.15s; }}
  #job-table tbody tr:hover {{ background:rgba(74,144,217,0.1); }}
  #job-table tbody td {{ padding:10px 14px; font-size:0.88rem; vertical-align:middle; }}

  .td-num      {{ color:#7a9ec0; font-weight:600; width:36px; }}
  .td-role     {{ color:#7ec8f7; font-weight:700; cursor:pointer; }}
  .td-role:hover {{ color:#b8e0ff; text-decoration:underline; }}
  .td-location {{ color:#b0c8e8; }}
  .td-date     {{ color:#8aaccc; white-space:nowrap; }}

  .apply-btn {{ display:inline-block;
    background:linear-gradient(90deg,#1660b0,#4a90d9);
    color:#fff; font-weight:700; font-size:0.78rem;
    padding:5px 14px; border-radius:6px; text-decoration:none;
    white-space:nowrap; transition:all 0.2s; }}
  .apply-btn:hover {{ background:linear-gradient(90deg,#4a90d9,#7ec8f7);
    transform:translateY(-1px); }}

  /* ── Modal overlay ───────────────────────────────────── */
  #modal-overlay {{
    display:none;
    position:fixed; inset:0;
    background:rgba(0,0,0,0.65);
    backdrop-filter:blur(4px);
    z-index:9998;
  }}

  /* ── Modal card — centered, never clips ──────────────── */
  #modal {{
    display:none;
    position:fixed;
    top:50%; left:50%;
    transform:translate(-50%,-50%);
    z-index:9999;
    width: min(420px, 92vw);
    max-height: 90vh;
    overflow-y: auto;
    background:linear-gradient(145deg,rgba(2,8,30,0.99),rgba(4,20,60,0.99));
    border:1px solid rgba(74,144,217,0.6);
    border-radius:18px;
    padding:22px 22px 18px;
    box-shadow:0 24px 70px rgba(0,0,0,0.9);
    color:#e8eef8;
    line-height:1.5;
    animation: popIn 0.18s ease;
  }}

  @keyframes popIn {{
    from {{ opacity:0; transform:translate(-50%,-50%) scale(0.93); }}
    to   {{ opacity:1; transform:translate(-50%,-50%) scale(1); }}
  }}

  #modal-close {{
    position:absolute; top:12px; right:14px;
    background:rgba(74,144,217,0.15);
    border:1px solid rgba(74,144,217,0.35);
    color:#7ec8f7; font-size:1.1rem; font-weight:700;
    width:28px; height:28px; border-radius:50%;
    cursor:pointer; display:flex; align-items:center; justify-content:center;
    transition:background 0.15s;
  }}
  #modal-close:hover {{ background:rgba(74,144,217,0.35); }}

  #modal-role {{
    font-size:1.05rem; font-weight:800; color:#fff;
    margin-bottom:12px; padding-right:30px; line-height:1.3;
  }}

  .m-div {{ border:none; border-top:1px solid rgba(74,144,217,0.2); margin:11px 0; }}

  .m-row {{
    display:flex; align-items:flex-start; justify-content:space-between;
    margin-bottom:10px;
  }}
  .m-label {{ font-size:0.7rem; color:#7ec8f7; text-transform:uppercase;
    font-weight:700; letter-spacing:0.5px; display:flex; align-items:center;
    gap:5px; padding-top:2px; }}
  .m-value {{ font-size:0.88rem; color:#ddeeff; font-weight:700; }}

  .m-rating-src {{ font-size:0.62rem; color:#556a88; margin-top:2px; }}
  .m-employees  {{ font-size:0.68rem; color:#7a9ec0; margin-top:2px; }}

  .m-location {{
    display:flex; align-items:flex-start; gap:7px;
    font-size:0.88rem; color:#b0cce8; margin-bottom:11px;
  }}
  .m-hq {{ display:block; font-size:0.68rem; color:#7a9ec0; margin-top:2px; }}

  .m-salary-box {{
    background:rgba(74,200,120,0.08);
    border:1px solid rgba(74,200,120,0.3);
    border-radius:11px; padding:12px 14px;
    margin-bottom:12px;
    display:flex; align-items:center; justify-content:space-between;
  }}
  .m-salary-label {{ font-size:0.68rem; color:#4ade80; text-transform:uppercase;
    font-weight:700; letter-spacing:0.5px; margin-bottom:4px; }}
  .m-salary-val {{ font-size:1.28rem; font-weight:900; color:#4ade80; }}
  .m-salary-sub {{ font-size:0.61rem; color:#5a9a72; margin-top:3px; line-height:1.4; }}

  .m-skills-box {{
    background:rgba(74,144,217,0.07);
    border:1px solid rgba(74,144,217,0.22);
    border-radius:11px; padding:12px 14px;
  }}
  .m-skills-label {{ font-size:0.72rem; color:#7ec8f7; text-transform:uppercase;
    font-weight:800; letter-spacing:0.6px; margin-bottom:9px; display:block; }}
  .skill-tag {{
    display:inline-block;
    background:rgba(74,144,217,0.18);
    border:1px solid rgba(74,144,217,0.45);
    border-radius:20px; padding:4px 12px;
    margin:3px 3px 3px 0;
    font-size:0.8rem; font-weight:700; color:#d4eeff;
  }}

  .m-hint {{ font-size:0.63rem; color:#3a4f66; text-align:center; margin-top:11px; }}
</style>
</head>
<body>

<div id="modal-overlay" onclick="closeModal()"></div>

<div id="modal">
  <div id="modal-close" onclick="closeModal()">X</div>
  <div id="modal-role"></div>
  <div class="m-location">
    <div><span id="m-location"></span>
      <span class="m-hq" id="m-hq"></span>
    </div>
  </div>
  <hr class="m-div">
  <div class="m-row">
    <span class="m-label">Rating</span>
    <div style="text-align:right">
      <span id="m-stars"></span>
      <div class="m-rating-src" id="m-rating-src"></div>
    </div>
  </div>
  <div class="m-row">
    <span class="m-label">Company Size</span>
    <div style="text-align:right">
      <span class="m-value" id="m-size"></span>
      <div class="m-employees" id="m-employees"></div>
    </div>
  </div>
  <hr class="m-div">
  <div class="m-salary-box">
    <div>
      <div class="m-salary-label">Internship Stipend — India</div>
      <div class="m-salary-val" id="m-salary"></div>
      <div class="m-salary-sub">Source: Glassdoor / AmbitionBox / Levels.fyi · ₹ INR</div>
    </div>
  </div>
  <div class="m-skills-box">
    <span class="m-skills-label">Core Skills Required</span>
    <div id="m-skills"></div>
  </div>
  <div class="m-hint">Press Esc or click outside to close</div>
</div>

<table id="job-table">
  <thead><tr>
    <th>#</th><th>Role</th><th>Location</th><th>Date</th><th>Apply</th>
  </tr></thead>
  <tbody id="job-body"></tbody>
</table>

<script>
(function(){{
  var data  = {rows_json};
  var tbody = document.getElementById('job-body');
  var modal = document.getElementById('modal');
  var overlay = document.getElementById('modal-overlay');

  data.forEach(function(row, idx){{
    var tr = document.createElement('tr');
    tr.innerHTML =
      '<td class="td-num">' + (idx+1) + '</td>'+
      '<td class="td-role" data-idx="'+idx+'">' + esc(row.role) + '</td>'+
      '<td class="td-location">' + esc(row.location) + '</td>'+
      '<td class="td-date">'     + esc(row.date)     + '</td>'+
      '<td><a class="apply-btn" href="'+row.url+'" target="_blank">Apply Link</a></td>';
    tbody.appendChild(tr);
  }});

  function stars(r){{
    var f=Math.floor(r), h=(r-f)>=0.5?1:0, e=5-f-h, s='';
    for(var i=0;i<f;i++) s+='<span style="color:#f5c518;font-size:1.1rem;">star</span>';
    for(var i=0;i<h;i++) s+='<span style="color:#f5c518;font-size:0.9rem;">½</span>';
    for(var i=0;i<e;i++) s+='<span style="color:#333;font-size:1.1rem;">star</span>';
    s+=' <span style="color:#aaa;font-size:0.8rem;">('+r.toFixed(1)+')</span>';
    return s;
  }}

  function openModal(idx){{
    var d=data[idx], n=d.intel;
    document.getElementById('modal-role').textContent  = d.role;
    document.getElementById('m-location').textContent  = d.location;
    document.getElementById('m-hq').textContent        = n.hq ? 'HQ: ' + n.hq : '';
    document.getElementById('m-size').textContent      = n.size;
    document.getElementById('m-employees').textContent = n.employees ? n.employees + ' employees' : '';
    document.getElementById('m-salary').textContent    = n.salary;

    var ratingEl = document.getElementById('m-stars');
    var srcEl    = document.getElementById('m-rating-src');
    if(n.rating === 0){{
      ratingEl.innerHTML = '<span style="color:#556a88;font-size:0.82rem;">N/A</span>';
      srcEl.textContent  = '';
    }} else {{
      ratingEl.innerHTML = stars(n.rating);
      srcEl.textContent  = 'Source: ' + n.rating_src;
    }}

    var sc = document.getElementById('m-skills');
    sc.innerHTML = '';
    n.requirements.forEach(function(r){{
      var sp = document.createElement('span');
      sp.className = 'skill-tag';
      sp.textContent = r;
      sc.appendChild(sp);
    }});

    overlay.style.display = 'block';
    modal.style.display   = 'block';
  }}

  window.closeModal = function(){{
    modal.style.display   = 'none';
    overlay.style.display = 'none';
  }};

  document.addEventListener('keydown', function(e){{
    if(e.key === 'Escape') closeModal();
  }});

  tbody.addEventListener('click', function(e){{
    var td = e.target.closest('.td-role');
    if(!td) return;
    openModal(parseInt(td.getAttribute('data-idx')));
  }});

  function esc(s){{
    return String(s).replace(/&/g,'&amp;').replace(/</g,'&lt;')
      .replace(/>/g,'&gt;').replace(/"/g,'&quot;');
  }}
}})();
</script>
</body></html>"""

    height = 56 + min(len(rows_data), NUM_POSTING) * 46 + 20
    st.components.v1.html(html, height=min(height, 2600), scrolling=True)

    # ── Prep for Interview buttons (one per matched job) ──────────────────
    st.markdown(
        "<p style='color:#9ab8d8;font-size:0.82rem;margin:12px 0 6px 0;'>"
        "<b>Interview Prep</b> — click a button below to generate questions for that role:</p>",
        unsafe_allow_html=True
    )
    cols = st.columns(min(len(rows_data), 4))
    for i, job in enumerate(rows_data):
        col_idx = i % 4
        with cols[col_idx]:
            btn_label = f"🎯 {job['role'][:28]}{'…' if len(job['role']) > 28 else ''}"
            if st.button(btn_label, key=f"prep_btn_{job['company']}_{job['role']}_{i}"):
                st.session_state["prep_company"] = job["company"]
                st.session_state["prep_role"]    = job["role"]
                st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# TECH SKILL VOCABULARY & ROLE PATTERNS
# ══════════════════════════════════════════════════════════════════════════════

TECH_SKILLS_VOCAB = [
    "machine learning","deep learning","natural language processing","computer vision",
    "data science","data analysis","data engineering","full stack","full-stack",
    "software engineer","software developer","web development","mobile development",
    "android development","ios development","embedded systems","system design",
    "devops engineer","site reliability","cloud computing","cybersecurity","blockchain",
    "python","javascript","typescript","java","c++","c#","golang","rust","kotlin","swift",
    "scala","ruby","php","dart","flutter","bash","matlab","r programming",
    "react","angular","vue","nodejs","django","flask","spring","fastapi","express",
    "tensorflow","pytorch","keras","scikit-learn","opencv","nextjs","nestjs","laravel",
    "mysql","postgresql","mongodb","redis","firebase","elasticsearch","dynamodb","sqlite",
    "aws","azure","gcp","docker","kubernetes","jenkins","terraform","ansible","linux",
    "ci/cd","github actions","gitlab","nginx",
    "power bi","tableau","excel","hadoop","spark","airflow","mlflow","pandas","numpy",
    "sql","jupyter","etl","databricks","snowflake",
    "api","rest api","graphql","microservices","iot","arduino","raspberry pi",
    "android","ios","react native","unity","unreal engine","solidity","web3",
]

ROLE_PATTERNS = [
    (["machine learning","deep learning","pytorch","tensorflow","nlp","computer vision","keras","neural network"],
     "Machine Learning Engineer", "machine-learning"),
    (["data science","data scientist","pandas","numpy","scikit-learn","sklearn","jupyter","matplotlib","seaborn","statistics"],
     "Data Science", "data-science"),
    (["data analyst","power bi","tableau","business analyst","excel","data analysis","google analytics"],
     "Data Analyst", "data-analytics"),
    (["android","kotlin","android studio","android sdk","mobile development"],
     "Mobile App Developer", "android-development"),
    (["ios","swift","xcode","swiftui","uikit"],
     "iOS Developer", "android-development"),
    (["flutter","dart","react native","mobile app"],
     "Mobile App Developer", "android-development"),
    (["devops","docker","kubernetes","jenkins","terraform","ansible","ci/cd","gitlab ci","github actions"],
     "DevOps Engineer", "devops"),
    (["cybersecurity","penetration testing","ethical hacking","soc","siem","kali linux","burp suite","nmap"],
     "Cybersecurity Engineer", "cybersecurity"),
    (["blockchain","solidity","web3","ethereum","smart contract","defi","truffle","hardhat"],
     "Blockchain Developer", "blockchain"),
    (["embedded","iot","rtos","firmware","arduino","raspberry pi","microcontroller","verilog","vhdl"],
     "Embedded Systems Engineer", "embedded-systems"),
    (["ui","ux","figma","adobe xd","sketch","prototyping","wireframe","user research","design system"],
     "UI/UX Designer", "graphic-design"),
    (["full stack","mern","mean","nextjs","nuxtjs","django","flask","spring boot","nodejs","express"],
     "Full Stack Developer", "web-development"),
    (["react","angular","vue","html","css","javascript","typescript","tailwind","bootstrap","frontend"],
     "Frontend Developer", "web-development"),
    (["spring","spring boot","microservices","rest api","graphql","backend","postgresql","mongodb","redis","kafka"],
     "Backend Developer", "web-development"),
    (["java","python","c++","algorithms","data structures","leetcode","competitive programming","oop","system design"],
     "Software Engineer", "software-development"),
]


def extract_smart_keywords(raw_text: str) -> dict:
    text_lower = raw_text.lower()
    found_skills = []
    for skill in TECH_SKILLS_VOCAB:
        if skill in text_lower and skill not in found_skills:
            found_skills.append(skill)
        if len(found_skills) >= 15:
            break

    best_role  = "Software Engineer"
    best_cat   = "software-development"
    best_score = 0
    for kw_list, role, cat in ROLE_PATTERNS:
        score = sum(1 for kw in kw_list if kw in text_lower)
        if score > best_score:
            best_score = score
            best_role  = role
            best_cat   = cat

    top5 = found_skills[:5] if found_skills else ["python", "java", "sql"]

    return {
        "role":            best_role,
        "internshala_cat": best_cat,
        "skills":          found_skills[:12],
        "top5":            top5,
    }


# ══════════════════════════════════════════════════════════════════════════════
# LIVE SEARCH PANEL
# ══════════════════════════════════════════════════════════════════════════════

def render_live_search(kw: dict):
    import urllib.parse as ul

    st.markdown("---")
    st.markdown("### Live Internship Search")
    st.markdown(
        f"<p style='color:#9ab8d8;font-size:0.85rem;margin:0 0 12px 0;'>"
        f"Detected role: <b style='color:#7ec8f7;'>{kw['role']}</b> &nbsp;·&nbsp; "
        f"Skills: <span style='color:#b0cce8;'>{', '.join(kw['skills'][:6]) or 'N/A'}</span></p>",
        unsafe_allow_html=True
    )

    with st.expander("Customise search query", expanded=True):
        c1, c2, c3 = st.columns([2, 2, 1])
        with c1:
            custom_role = st.text_input("Role / Title", value=kw["role"],
                placeholder="e.g. Data Scientist, React Developer", key="custom_role")
        with c2:
            custom_skills = st.text_input("Skills (comma-separated)",
                value=", ".join(kw["top5"][:4]),
                placeholder="e.g. Python, React, SQL", key="custom_skills")
        with c3:
            custom_location = st.selectbox("Location",
                ["India","Bangalore","Mumbai","Hyderabad","Pune",
                 "Chennai","Delhi","Remote","Noida","Gurugram"],
                key="custom_location")

        c4, c5 = st.columns([2, 2])
        with c4:
            custom_time = st.selectbox("Date Posted",
                ["Last 7 days","Last 24 hours","Last 30 days","Any time"],
                key="custom_time")
        with c5:
            custom_type = st.selectbox("Opportunity Type",
                ["Internship","Internship + Full-time","Full-time only"],
                key="custom_type")

        skills_list = [s.strip() for s in custom_skills.split(",") if s.strip()]
        preview_q = f"{custom_role} intern"
        if skills_list:
            preview_q += f" {skills_list[0]}"

        st.markdown(
            f"<div style='background:rgba(74,144,217,0.08);border:1px solid rgba(74,144,217,0.25);"
            f"border-radius:8px;padding:10px 14px;margin-top:8px;"
            f"display:flex;align-items:center;gap:10px;flex-wrap:wrap;'>"
            f"<span style='color:#556a88;font-size:0.75rem;font-weight:700;'>SEARCH PREVIEW</span>"
            f"<code style='background:rgba(74,144,217,0.15);padding:3px 10px;border-radius:5px;"
            f"color:#7ec8f7;font-size:0.82rem;font-weight:700;'>{preview_q}</code>"
            f"<span style='color:#445566;font-size:0.72rem;'>in {custom_location} · {custom_time}</span>"
            f"</div>",
            unsafe_allow_html=True
        )

    role      = custom_role.strip() or kw["role"]
    top5      = skills_list[:5] if skills_list else kw["top5"]
    location  = custom_location

    time_map = {
        "Last 24 hours": {"li": "r86400",   "naukri": "1",  "indeed": "1",  "tbs": "qdr:d"},
        "Last 7 days":   {"li": "r604800",  "naukri": "7",  "indeed": "7",  "tbs": "qdr:w"},
        "Last 30 days":  {"li": "r2592000", "naukri": "30", "indeed": "30", "tbs": "qdr:m"},
        "Any time":      {"li": "",          "naukri": "",   "indeed": "",   "tbs": ""},
    }
    tf = time_map.get(custom_time, time_map["Last 7 days"])

    jt_map = {
        "Internship":             "I",
        "Internship + Full-time": "I%2CF",
        "Full-time only":         "F",
    }
    li_jt = jt_map.get(custom_type, "I")

    role_q    = ul.quote(role)
    role_plus = role.replace(" ", "+")
    role_slug = role.lower().replace(" ", "-")
    loc_q     = ul.quote(location)
    skill0    = top5[0] if top5 else "python"
    skills_str = " ".join(top5[:3])
    combo_q   = ul.quote(f"{role} intern {skills_str}")

    li_time  = f"&f_TPR={tf['li']}"     if tf["li"]      else ""
    nk_age   = f"&jobAge={tf['naukri']}" if tf["naukri"]  else ""
    id_age   = f"&fromage={tf['indeed']}" if tf["indeed"] else ""
    g_tbs    = f"&tbs={tf['tbs']}"       if tf["tbs"]     else ""

    is_cat      = kw.get("internshala_cat", "software-development")
    skill0_slug = skill0.replace(" ", "-")

    platforms = [
        {"name":"LinkedIn Jobs","icon":"in","bg":"#0077B5",
         "url":(f"https://www.linkedin.com/jobs/search/?keywords={role_plus}+intern"
                f"&location={loc_q}&f_JT={li_jt}{li_time}&sortBy=DD"),
         "badge":"Updated daily","desc":f'"{role} intern" · {location} · {custom_time}'},
        {"name":"Internshala","icon":"IS","bg":"#00aaff",
         "url":f"https://internshala.com/internships/{is_cat}-internship/",
         "badge":"India #1","desc":f"{role} internships · freshest listings"},
        {"name":"Naukri.com","icon":"NK","bg":"#f06024",
         "url":(f"https://www.naukri.com/{role_slug}-internship-jobs"
                f"?typeId=1&src=jobsearchDesk{nk_age}"),
         "badge":"Freshers welcome","desc":f"{role} intern · {location} · {custom_time}"},
        {"name":"Indeed India","icon":"ID","bg":"#2557a7",
         "url":f"https://in.indeed.com/jobs?q={role_q}+internship&l={loc_q}&sort=date{id_age}",
         "badge":"Real-time","desc":f"Sorted by most recent · {custom_time}"},
        {"name":"Unstop","icon":"UN","bg":"#6c3de0",
         "url":f"https://unstop.com/internships?search={role_q}&oppstatus=open&sort=NEWEST",
         "badge":"Competitions+","desc":"Internships + hackathons · newest first"},
        {"name":"Glassdoor","icon":"GD","bg":"#0caa41",
         "url":(f"https://www.glassdoor.co.in/Job/{loc_q.lower()}-"
                f"{role_slug}-intern-jobs-SRCH_IL.0,{len(location)}"
                f"_IN115_KO{len(location)+1},{len(location)+1+len(role)+7}.htm?sortBy=date_desc"),
         "badge":"With salaries","desc":f"Ratings + salary insights · {location}"},
        {"name":"Wellfound","icon":"WF","bg":"#e07b39",
         "url":(f"https://wellfound.com/jobs?role={role_q}"
                f"&jobType=internship&locations%5B%5D={loc_q}"),
         "badge":"Startups only","desc":"VC-backed startups · equity + stipend"},
        {"name":"Google Jobs","icon":"G","bg":"#EA4335",
         "url":(f"https://www.google.com/search?q={combo_q}+{loc_q}+internship"
                f"{g_tbs}&ibp=htl;jobs"),
         "badge":"Aggregated","desc":f"All sites combined · {custom_time}"},
        {"name":"TimesJobs","icon":"TJ","bg":"#c0392b",
         "url":(f"https://www.timesjobs.com/candidate/internship-jobs.html"
                f"?searchType=personalizedSearch&from=submit"
                f"&txtKeywords={role_q}&txtLocation={loc_q}"),
         "badge":"Indian market","desc":f"{role} internships · {location}"},
        {"name":"Shine.com","icon":"SH","bg":"#1b3a6b",
         "url":f"https://www.shine.com/job-search/{role_slug}-internship-jobs-in-india",
         "badge":"Freshers","desc":"Entry-level + intern listings · India"},
        {"name":"HackerEarth","icon":"HE","bg":"#323754",
         "url":(f"https://www.hackerearth.com/challenges/hiring/"
                f"?skill={ul.quote(skill0)}&opportunity_type=internship"),
         "badge":"Skill-based","desc":f'Matched to skill: "{skill0}"'},
        {"name":"Internshala (Skills)","icon":"IS","bg":"#0088cc",
         "url":f"https://internshala.com/internships/keywords-{skill0_slug}-internship/",
         "badge":"Skill match","desc":f'Filtered by skill: "{skill0}"'},
    ]

    cards_html = """<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
  * { box-sizing:border-box; margin:0; padding:0; }
  body { background:transparent; font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif; }
  .grid { display:grid; grid-template-columns:repeat(4,1fr); gap:10px; padding:4px 2px; }
  .card {
    background:rgba(5,12,35,0.88); border:1px solid rgba(74,144,217,0.22);
    border-radius:12px; padding:13px 12px 11px;
    text-decoration:none; display:flex; flex-direction:column; gap:7px;
    transition:all 0.2s ease; position:relative; overflow:hidden;
  }
  .card::before { content:''; position:absolute; top:0; left:0; right:0;
    height:3px; background:var(--accent); opacity:0.85; }
  .card:hover { background:rgba(10,22,60,0.97); border-color:var(--accent);
    transform:translateY(-3px);
    box-shadow:0 8px 24px rgba(0,0,0,0.5),0 0 14px rgba(74,144,217,0.12); }
  .card-top { display:flex; align-items:center; gap:9px; }
  .avatar { width:34px; height:34px; border-radius:8px; background:var(--accent);
    display:flex; align-items:center; justify-content:center;
    font-size:0.67rem; font-weight:900; color:#fff; letter-spacing:-0.5px; flex-shrink:0; }
  .name { font-size:0.85rem; font-weight:800; color:#e8eef8; line-height:1.2; }
  .badge { display:inline-block; background:var(--accent); opacity:0.9;
    border-radius:4px; padding:1px 7px; font-size:0.59rem; font-weight:700;
    color:#fff; letter-spacing:0.3px; text-transform:uppercase; margin-top:2px; }
  .desc { font-size:0.71rem; color:#7a9ec0; line-height:1.45; flex-grow:1; }
  .arrow { position:absolute; top:11px; right:11px; font-size:0.88rem;
    color:rgba(74,144,217,0.35); }
  .card:hover .arrow { color:var(--accent); }
  .open-btn { display:block; background:var(--accent); opacity:0.88; color:#fff;
    font-size:0.69rem; font-weight:700; text-align:center;
    padding:5px 0; border-radius:6px; margin-top:auto; letter-spacing:0.3px; }
  .card:hover .open-btn { opacity:1; }
</style></head><body>
<div class="grid">
"""
    for p in platforms:
        cards_html += (
            f'  <a class="card" href="{p["url"]}" target="_blank" rel="noopener noreferrer" '
            f'style="--accent:{p["bg"]};">\n'
            f'    <span class="arrow">↗</span>\n'
            f'    <div class="card-top">\n'
            f'      <div class="avatar">{p["icon"]}</div>\n'
            f'      <div><div class="name">{p["name"]}</div>'
            f'<span class="badge">{p["badge"]}</span></div>\n'
            f'    </div>\n'
            f'    <div class="desc">{p["desc"]}</div>\n'
            f'    <span class="open-btn">Open Live Listings →</span>\n'
            f'  </a>\n'
        )
    cards_html += "</div></body></html>"

    st.components.v1.html(cards_html, height=430, scrolling=False)

    st.markdown(
        f"<div style='background:rgba(0,0,0,0.3);border:1px solid rgba(74,144,217,0.15);"
        f"border-radius:8px;padding:9px 14px;margin-top:6px;"
        f"display:flex;align-items:center;gap:10px;flex-wrap:wrap;'>"
        f"<span style='color:#445566;font-size:0.73rem;font-weight:600;'>Active query</span>"
        f"<code style='background:rgba(74,144,217,0.12);padding:2px 9px;border-radius:5px;"
        f"color:#7ec8f7;font-size:0.8rem;'>{role} intern</code>"
        f"<span style='color:#445566;font-size:0.71rem;'>"
        f"Skills: <span style='color:#b0cce8;'>{', '.join(top5[:3])}</span>"
        f" &nbsp;·&nbsp; {location} &nbsp;·&nbsp; {custom_time}</span>"
        f"</div>",
        unsafe_allow_html=True
    )
    st.markdown("---")


# ══════════════════════════════════════════════════════════════════════════════
# RESUME CUSTOMISER PAGE
# ══════════════════════════════════════════════════════════════════════════════

def _call_claude(system_prompt: str, user_prompt: str, max_tokens: int = 2000) -> str:
    """Call Claude API via Anthropic."""
    import urllib.request, json as _json
    body = _json.dumps({
        "model": "claude-sonnet-4-20250514",
        "max_tokens": max_tokens,
        "system": system_prompt,
        "messages": [{"role": "user", "content": user_prompt}]
    }).encode()
    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=body,
        headers={
            "Content-Type":      "application/json",
            "anthropic-version": "2023-06-01",
        },
        method="POST"
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            data = _json.loads(resp.read())
            return data["content"][0]["text"]
    except Exception as e:
        return f"[API Error: {e}]"


def page_customiser():
    """Full-page resume customiser — tailor resume to any internship."""

    if st.button("← Back to Recommendations", key="back_btn"):
        st.session_state["page"] = "home"
        st.rerun()

    st.markdown(
        "<h1 style='margin:0 0 4px 0;'>Resume Customiser</h1>"
        "<p style='color:#9ab8d8;font-size:0.9rem;margin-bottom:20px;'>"
        "Tailor your resume to match any internship job description — powered by Claude AI.</p>",
        unsafe_allow_html=True
    )

    left, right = st.columns([1, 1], gap="large")

    with left:
        st.markdown("#### Your Resume")
        prefill_resume = st.session_state.get("resume_raw", "")

        tab_upload, tab_paste = st.tabs(["Upload Resume", "Paste Text"])
        with tab_upload:
            resume_file = st.file_uploader(
                "Upload your resume (PDF / DOCX / TXT)",
                type=["pdf","docx","txt"],
                key="cust_upload"
            )
            if resume_file:
                prefill_resume = read_resume(resume_file)
                st.session_state["resume_raw"] = prefill_resume
                st.success(f"Loaded: {resume_file.name}")

        with tab_paste:
            pasted = st.text_area(
                "Or paste your resume text here",
                value=prefill_resume[:3000] if prefill_resume else "",
                height=200,
                key="cust_paste",
                placeholder="Paste the full text of your resume here..."
            )
            if pasted.strip():
                prefill_resume = pasted

        resume_text = prefill_resume.strip()
        if resume_text:
            wc = len(resume_text.split())
            st.caption(f"{wc} words detected in resume")

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("#### Internship Details")

        col_c, col_r = st.columns(2)
        with col_c:
            jd_company = st.text_input("Company name",
                placeholder="e.g. Google India, Swiggy…", key="jd_company")
        with col_r:
            jd_role = st.text_input("Role / Title",
                placeholder="e.g. Software Engineering Intern", key="jd_role")

        jd_text = st.text_area(
            "Job Description / Skills Required",
            height=200,
            key="jd_text",
            placeholder=(
                "Paste the full JD here, OR just list key skills/requirements:\n\n"
                "• Python, Machine Learning, TensorFlow\n"
                "• Experience with REST APIs\n"
                "• Strong DSA skills\n"
                "• Excellent communication"
            )
        )

        if jd_company and jd_role:
            try:
                from company_intelligence import get_company_info
                intel = get_company_info(jd_company, jd_role)
                auto_skills = "\n".join(f"• {s}" for s in intel["requirements"])
                if not jd_text.strip():
                    st.info(f"Auto-detected skills for **{jd_role}** at **{jd_company}**")
                    st.code(auto_skills, language="")
            except Exception:
                pass

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("#### Customisation Options")

        cust_opts = st.multiselect(
            "What should Claude improve?",
            options=[
                "Professional Summary / Objective",
                "Work Experience bullet points",
                "Projects section",
                "Skills section",
                "Education highlights",
                "Keywords for ATS (Applicant Tracking Systems)",
                "Action verbs & impact metrics",
                "Remove irrelevant content",
            ],
            default=[
                "Professional Summary / Objective",
                "Work Experience bullet points",
                "Skills section",
                "Keywords for ATS (Applicant Tracking Systems)",
            ],
            key="cust_opts"
        )

        tone = st.select_slider(
            "Writing tone",
            options=["Conservative", "Professional", "Confident", "Bold & Impactful"],
            value="Professional",
            key="cust_tone"
        )

        col_btn1, col_btn2 = st.columns([3, 1])
        with col_btn1:
            run_btn = st.button("Generate Tailored Resume", key="run_customiser",
                                use_container_width=True, type="primary")
        with col_btn2:
            clear_btn = st.button("Clear", key="clear_result", use_container_width=True)
        if clear_btn:
            for k in ["cust_result", "cust_analysis", "cust_missing"]:
                st.session_state.pop(k, None)
            st.rerun()

    with right:
        st.markdown("#### Tailored Resume Output")

        if run_btn:
            if not resume_text:
                st.error("Please provide your resume text first (upload or paste).")
            elif not jd_text.strip() and not jd_role.strip():
                st.error("Please enter the job description or role.")
            else:
                target   = f"{jd_role} at {jd_company}" if jd_company else jd_role or "the internship"
                opts_str = "\n".join(f"- {o}" for o in cust_opts) if cust_opts else "- All sections"

                system = (
                    "You are an expert resume writer who specialises in tailoring resumes "
                    "for internship applications in India's tech industry. "
                    "You write concise, ATS-optimised, impactful resumes. "
                    "Always preserve the candidate's real experience — never fabricate. "
                    "Use strong action verbs and quantify achievements where possible."
                )

                analysis_prompt = f"""Analyse this resume against the job description and provide:

1. MATCH SCORE (0-100): How well does the resume currently match?
2. MISSING KEYWORDS: List keywords in the JD not found in the resume (max 10)
3. STRENGTHS: 3 things that already align well
4. GAPS: 3 most critical gaps to address

RESUME:
{resume_text[:2500]}

TARGET ROLE: {target}
JOB DESCRIPTION / SKILLS:
{jd_text[:1500]}

Format your response as:
MATCH_SCORE: [number]
MISSING: [comma-separated keywords]
STRENGTHS: [bullet list]
GAPS: [bullet list]"""

                rewrite_prompt = f"""Rewrite and optimise this resume specifically for: {target}

ORIGINAL RESUME:
{resume_text[:3000]}

JOB DESCRIPTION / REQUIRED SKILLS:
{jd_text[:1500]}

SECTIONS TO IMPROVE:
{opts_str}

TONE: {tone}

INSTRUCTIONS:
- Rewrite the Professional Summary to directly address this role
- Rephrase experience/project bullets to highlight relevant skills from the JD
- Reorganise the Skills section to put the most relevant skills first
- Insert missing keywords from the JD naturally (don't stuff)
- Add ATS-friendly keywords throughout
- Keep ALL real experience — do not fabricate anything
- Use strong action verbs: Built, Developed, Optimised, Designed, Led, Improved...
- Add impact metrics where implied (e.g. "reduced load time by ~30%")
- Format clearly with section headers in CAPS

Return the complete tailored resume text, ready to copy into a document."""

                with st.spinner("Analysing your resume…"):
                    analysis = _call_claude(system, analysis_prompt, 600)
                    st.session_state["cust_analysis"] = analysis

                with st.spinner("Rewriting and tailoring your resume…"):
                    result = _call_claude(system, rewrite_prompt, 2000)
                    st.session_state["cust_result"] = result

                try:
                    score_match   = re.search(r"MATCH_SCORE:\s*(\d+)", analysis)
                    missing_match = re.search(r"MISSING:\s*(.+?)(?:\n|$)", analysis)
                    st.session_state["cust_score"]   = int(score_match.group(1)) if score_match else None
                    st.session_state["cust_missing"] = missing_match.group(1).strip() if missing_match else ""
                except Exception:
                    pass

        if "cust_analysis" in st.session_state:
            analysis = st.session_state["cust_analysis"]
            score    = st.session_state.get("cust_score")
            missing  = st.session_state.get("cust_missing", "")

            if score is not None:
                color = "#ef4444" if score < 50 else "#f59e0b" if score < 75 else "#22c55e"
                label = "Needs Work" if score < 50 else "Good Match" if score < 75 else "Strong Match"
                missing_html = (
                    f"<div style='color:#f59e0b;font-size:0.72rem;margin-top:4px;'>"
                    f"Missing keywords: <b style='color:#fcd34d;'>{missing}</b></div>"
                    if missing else ""
                )
                st.markdown(
                    f"<div style='display:flex;align-items:center;gap:14px;"
                    f"background:rgba(0,0,0,0.3);border:1px solid {color}44;"
                    f"border-radius:10px;padding:12px 16px;margin-bottom:12px;'>"
                    f"<div style='text-align:center;'>"
                    f"<div style='font-size:2rem;font-weight:900;color:{color};'>{score}</div>"
                    f"<div style='font-size:0.65rem;color:{color};font-weight:700;"
                    f"text-transform:uppercase;'>{label}</div></div>"
                    f"<div>"
                    f"<div style='color:#e8eef8;font-size:0.85rem;font-weight:700;'>"
                    f"Resume Match Score</div>"
                    f"<div style='color:#7a9ec0;font-size:0.75rem;'>"
                    f"Before tailoring — Claude will improve this</div>"
                    f"{missing_html}"
                    f"</div></div>",
                    unsafe_allow_html=True
                )

            with st.expander("Full Gap Analysis", expanded=False):
                st.text(analysis)

        if "cust_result" in st.session_state:
            result = st.session_state["cust_result"]

            st.markdown(
                "<div style='background:rgba(74,200,120,0.08);border:1px solid rgba(74,200,120,0.3);"
                "border-radius:8px;padding:8px 14px;margin-bottom:10px;"
                "display:flex;align-items:center;gap:8px;'>"
                "<span style='color:#4ade80;font-size:1rem;font-weight:700;'>&#10003;</span>"
                "<span style='color:#4ade80;font-size:0.82rem;font-weight:700;'>"
                "Tailored resume generated! Copy it below or download as TXT.</span></div>",
                unsafe_allow_html=True
            )

            st.text_area("Your tailored resume (copy this):",
                         value=result, height=450, key="result_display")

            col_d1, col_d2 = st.columns(2)
            company_slug = (jd_company or "internship").lower().replace(" ","_")
            role_slug    = (jd_role    or "role").lower().replace(" ","_")
            filename     = f"resume_{company_slug}_{role_slug}.txt"

            with col_d1:
                st.download_button("Download as TXT", data=result.encode("utf-8"),
                    file_name=filename, mime="text/plain",
                    use_container_width=True, key="dl_txt")
            with col_d2:
                st.download_button("Download as Markdown", data=result.encode("utf-8"),
                    file_name=filename.replace(".txt",".md"), mime="text/markdown",
                    use_container_width=True, key="dl_md")

            st.markdown(
                "<div style='background:rgba(74,144,217,0.08);"
                "border:1px solid rgba(74,144,217,0.2);"
                "border-radius:8px;padding:12px 14px;margin-top:12px;'>"
                "<div style='color:#7ec8f7;font-size:0.72rem;font-weight:700;"
                "text-transform:uppercase;letter-spacing:0.5px;margin-bottom:8px;'>"
                "Next Steps</div>"
                "<ul style='color:#9ab8d8;font-size:0.78rem;margin:0;padding-left:1.1rem;"
                "line-height:1.8;'>"
                "<li>Copy tailored resume → paste into <b>Canva / Overleaf / Google Docs</b></li>"
                "<li>Review every bullet — ensure all facts are accurate</li>"
                "<li>Run through <b>Jobscan.co</b> for ATS score verification</li>"
                "<li>Tailor further for each company using the JD keywords highlighted above</li>"
                "</ul></div>",
                unsafe_allow_html=True
            )

        elif "cust_result" not in st.session_state:
            st.markdown(
                "<div style='background:rgba(0,0,0,0.25);border:2px dashed rgba(74,144,217,0.2);"
                "border-radius:12px;padding:48px 24px;text-align:center;margin-top:8px;'>"
                "<div style='font-size:2rem;margin-bottom:12px;color:#2a4a6a;font-weight:700;'>—</div>"
                "<div style='color:#4a6080;font-size:0.95rem;font-weight:600;'>"
                "Your tailored resume will appear here</div>"
                "<div style='color:#2a3a50;font-size:0.78rem;margin-top:6px;'>"
                "Fill in your resume + JD on the left, then click Generate</div>"
                "</div>",
                unsafe_allow_html=True
            )


# ══════════════════════════════════════════════════════════════════════════════
# INTERVIEW PREP PAGE  (uses local Ollama — no internet required)
# ══════════════════════════════════════════════════════════════════════════════

def render_interview_prep(company: str, role: str):
    """Render interview prep for a specific role using local Ollama LLM."""
    st.markdown(f"### Interview Prep — {role} @ {company}")

    if st.button("← Back"):
        st.session_state.pop("prep_company", None)
        st.session_state.pop("prep_role", None)
        st.rerun()

    resume_text = st.session_state.get("resume_raw", "")
    cache_key   = f"prep_{company}_{role}"

    if cache_key not in st.session_state:
        prompt = f"""You are a senior technical interviewer at an Indian tech company.

Generate interview prep for:
Role: {role} at {company}

Return exactly this JSON format only, with no markdown and no extra explanation:
{{
  "technical": ["Q1","Q2","Q3","Q4","Q5"],
  "behavioral": ["Q1","Q2","Q3"],
  "company_specific": ["Q1","Q2"],
  "tips": ["tip1","tip2","tip3"],
  "resources": ["resource1","resource2"]
}}

Guidelines:
- Technical questions should match the role.
- Behavioral questions should be internship-friendly.
- Company-specific questions should be based on the role and company type.
- Tips should be short and practical.
- Resources should be concise.

{"Resume context: " + resume_text[:500] if resume_text else ""}
"""

        with st.spinner("Generating interview questions via Ollama (phi3:mini)…"):
            raw = _call_local_llm(prompt)

        try:
            clean = re.sub(r"```(?:json)?", "", raw).strip().strip("`")
            data  = json.loads(clean)

            if not isinstance(data, dict):
                raise ValueError("Invalid JSON structure")

            data.setdefault("technical",       [])
            data.setdefault("behavioral",      [])
            data.setdefault("company_specific",[])
            data.setdefault("tips",            [])
            data.setdefault("resources",       [])

        except Exception:
            data = {
                "technical":        [raw] if raw else ["Could not generate technical questions."],
                "behavioral":       [],
                "company_specific": [],
                "tips":             [],
                "resources":        []
            }

        st.session_state[cache_key] = data

    data = st.session_state[cache_key]

    t1, t2, t3, t4 = st.tabs(["Technical", "Behavioral", "Company", "Tips & Resources"])

    with t1:
        technical_questions = data.get("technical", [])
        if technical_questions:
            for i, q in enumerate(technical_questions, 1):
                with st.expander(f"Q{i}. {q}"):
                    st.text_area("Your answer (practice):",
                                 key=f"ans_{cache_key}_t_{i}", height=120)
        else:
            st.info("No technical questions generated.")

    with t2:
        behavioral_questions = data.get("behavioral", [])
        if behavioral_questions:
            for i, q in enumerate(behavioral_questions, 1):
                with st.expander(f"Q{i}. {q}"):
                    st.text_area("STAR format answer:",
                                 key=f"ans_{cache_key}_b_{i}", height=120)
        else:
            st.info("No behavioral questions generated.")

    with t3:
        company_questions = data.get("company_specific", [])
        if company_questions:
            for q in company_questions:
                st.markdown(f"- {q}")
        else:
            st.info("No company-specific questions generated.")

    with t4:
        tips      = data.get("tips", [])
        resources = data.get("resources", [])

        if tips:
            st.markdown("**Tips:**")
            for tip in tips:
                st.markdown(f"- {tip}")
        else:
            st.info("No tips generated.")

        if resources:
            st.markdown("**Resources:**")
            for r in resources:
                st.markdown(f"- {r}")
        else:
            st.info("No resources generated.")


# ══════════════════════════════════════════════════════════════════════════════
# NLP HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def keep_alpha_char(text):
    return re.sub(r'\s+', ' ', re.sub(r'[^a-zA-Z]', ' ', text))

def nltk_pos_tagger(tag):
    if tag.startswith('J'):   return wordnet.ADJ
    elif tag.startswith('V'): return wordnet.VERB
    elif tag.startswith('N'): return wordnet.NOUN
    elif tag.startswith('R'): return wordnet.ADV
    else:                     return None

def lemmatize_sentence(sentence):
    lem    = WordNetLemmatizer()
    tagged = nltk.pos_tag(nltk.word_tokenize(sentence))
    result = []
    for word, tag in tagged:
        wn = nltk_pos_tagger(tag)
        if wn:
            result.append(lem.lemmatize(word, wn))
    return " ".join(result)

def remove_stop_words(text):
    stops = set(stopwords.words('english'))
    return ' '.join(w for w in nltk.word_tokenize(str(text)) if w.lower() not in stops)

def pre_process_resume(text):
    return remove_stop_words(lemmatize_sentence(keep_alpha_char(text))).lower()

def pre_process_data_job(df):
    df.dropna(subset=['Role'], inplace=True)
    df['data'] = (df['Role'].apply(keep_alpha_char)
                            .apply(lemmatize_sentence)
                            .apply(remove_stop_words)
                            .str.lower())
    return df

def recommend_job(resume_text, mat, vec, df):
    rvec = vec.transform([resume_text])
    sims = cosine_similarity(rvec, mat)
    return pd.DataFrame([df.iloc[i] for i in sims.argsort()[0][::-1]])

def return_table_job(resume_text, job_df):
    vec  = TfidfVectorizer(stop_words='english')
    mat  = vec.fit_transform(job_df['data'])
    recs = recommend_job(resume_text, mat, vec, job_df)
    return recs.drop(columns=recs.columns[0])

def get_top_features(resume_text, job_df):
    vec   = TfidfVectorizer(stop_words='english')
    mat   = vec.fit_transform(job_df['data'])
    rvec  = vec.transform([resume_text])
    names = np.array(vec.get_feature_names_out())
    return names[np.argsort(rvec.toarray()[0])[::-1]]

def get_job_df():
    rows = return_data_list()
    df   = pd.DataFrame(rows, columns=['Company','Role','Location','Application/Link','Date Posted'])
    return pre_process_data_job(df)

def post_process_table(resume_text, job_df):
    df = return_table_job(str(resume_text), job_df)
    df = df.head(NUM_POSTING).sort_index(ascending=True)
    df = df.reset_index(drop=True).iloc[:, 0:-1]
    df.index = df.index + 1
    return df

def return_data_list():
    with open('WebCrawler/table.html', 'r', encoding='utf-8') as file:
        html_content = file.read()
    soup         = BeautifulSoup(html_content, 'html.parser')
    table        = soup.find('table')
    rows         = []
    last_company = ""
    for tr in table.find_all('tr'):
        tds = tr.find_all('td')
        if not tds or len(tds) < 2:
            continue
        row = [
            td.text.strip() if td.text.strip() else (td.find('a').get('href') if td.find('a') else '')
            for td in tds
        ]
        if row[0].strip() in ('↳', ''):
            row[0] = last_company
        else:
            last_company = row[0]
        rows.append(row)
    return rows

def download_csv(df):
    st.markdown("**Download matched internship postings**")
    csv_buffer   = BytesIO()
    excel_buffer = BytesIO()
    df.to_csv(csv_buffer, index=False)
    col1, col2 = st.columns(2)
    with col1:
        st.download_button("Download CSV",   csv_buffer.getvalue(),
                           "internship_posting.csv", "text/csv")
    with col2:
        st.download_button("Download Excel", excel_buffer.getvalue(),
                           "internship_posting.xlsx", "text/xlsx")

def display_features_slider(resume, job_df):
    st.markdown("### Resume Keywords")
    n = st.slider("Number of top keywords to display:", min_value=1, max_value=100, value=5)
    top_features = get_top_features(resume, job_df)
    st.markdown(f"**Top {n} keywords from your resume:**")
    st.table(pd.DataFrame({"Keywords": top_features[:n]}))


# ══════════════════════════════════════════════════════════════════════════════
# MAIN — PAGE ROUTER
# ══════════════════════════════════════════════════════════════════════════════

def main():
    st.set_page_config(page_title="Internship Recommender", layout="wide")

    # ── Route 1: Interview Prep page ──────────────────────────────────────
    if "prep_company" in st.session_state and "prep_role" in st.session_state:
        render_interview_prep(
            st.session_state["prep_company"],
            st.session_state["prep_role"]
        )
        return  # ← INSIDE the if-block (correct indentation)

    # ── Route 2: Resume Customiser page ───────────────────────────────────
    if "page" not in st.session_state:
        st.session_state["page"] = "home"

    if st.session_state["page"] == "customiser":
        page_customiser()
        return

    # ── Route 3: Home page ────────────────────────────────────────────────
    nltk.download('punkt',                      quiet=True)
    nltk.download('averaged_perceptron_tagger', quiet=True)
    nltk.download('wordnet',                    quiet=True)
    nltk.download('stopwords',                  quiet=True)
    nltk.download('punkt_tab',                  quiet=True)

    with open('style.css') as f:
        css = f.read()

    for ext, mime in [("jpg","image/jpeg"),("jpeg","image/jpeg"),("png","image/png")]:
        try:
            b64 = get_base64_image(f"image-1.{ext}")
            css = css.replace('url("image-1.jpg")', f'url("data:{mime};base64,{b64}")')
            break
        except FileNotFoundError:
            continue

    st.markdown(f'<style>{css}</style>', unsafe_allow_html=True)
    st.markdown(
        "<h1 style='margin:0 0 2px 0;font-size:1.7rem;font-weight:900;color:#e8eef8;'>"
        "Internship Recommender</h1>"
        "<p style='color:#7a9ec0;font-size:0.88rem;margin-bottom:20px;'>"
        "Upload your resume — we match it against live internship postings using NLP. "
        "Click any role for salary, skills and company details.</p>",
        unsafe_allow_html=True
    )

    uploaded_file = st.file_uploader(
        "Upload your Resume — PDF, DOCX, or TXT",
        type=["pdf", "docx", "txt"]
    )

    if uploaded_file is not None:
        st.success(f"Uploaded: {uploaded_file.name}")
        st.caption(f"Format: {uploaded_file.name.split('.')[-1].upper()}  ·  "
                   f"{round(uploaded_file.size / 1024, 1)} KB")

        raw_text = read_resume(uploaded_file)
        if not raw_text:
            st.error("Could not extract text from your file.")
            return

        # Store raw text so interview prep & customiser can access it
        st.session_state["resume_raw"] = raw_text

        with st.expander("Preview extracted resume text"):
            st.text(raw_text[:2000] + ("..." if len(raw_text) > 2000 else ""))

        with st.spinner("Analysing your resume and matching internships…"):
            resume           = pre_process_resume(raw_text)
            job_df           = get_job_df()
            df_resume_sorted = post_process_table(resume, job_df)

        # Live job search section
        smart_keywords = extract_smart_keywords(raw_text)
        render_live_search(smart_keywords)

        display_features_slider(resume, job_df)
        download_csv(df_resume_sorted)

        st.markdown("### Top Internship Matches")
        st.markdown(
            "<p style='color:#9ab8d8;font-size:0.82rem;margin-bottom:0.6rem;'>"
            "Click any <strong>Role</strong> to see a full details card — stipend, skills, rating.<br>"
            "Use the <strong>🎯 buttons below the table</strong> to launch Interview Prep for any role.</p>",
            unsafe_allow_html=True
        )

        # Table + prep buttons (buttons are rendered inside render_table_with_tooltips)
        render_table_with_tooltips(df_resume_sorted)


if __name__ == "__main__":
    main()