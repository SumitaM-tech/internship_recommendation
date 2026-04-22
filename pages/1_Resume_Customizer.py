"""pages/1_Resume_Customizer.py — AI Resume Builder"""
import json
import re
import copy
import sys
import os
import urllib.request
import urllib.error
from io import BytesIO

import streamlit as st

_parent = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if _parent not in sys.path:
    sys.path.insert(0, _parent)

try:
    from ats_scorer import score_resume, render_ats_card
    _ATS_AVAILABLE = True
except ImportError:
    _ATS_AVAILABLE = False

from PyPDF2 import PdfReader
import docx
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

OLLAMA_BASE   = "http://localhost:11434"
DEFAULT_MODEL = "phi3:mini"

def _get_model() -> str:
    return st.session_state.get("ollama_model_pick",
        st.session_state.get("ollama_model_manual", DEFAULT_MODEL))

def _ollama_available() -> tuple:
    try:
        req = urllib.request.Request(f"{OLLAMA_BASE}/api/tags", method="GET")
        with urllib.request.urlopen(req, timeout=3) as r:
            data = json.loads(r.read())
        return True, [m["name"] for m in data.get("models", [])], None
    except Exception as e:
        return False, [], str(e)

st.set_page_config(page_title="Resume Builder", page_icon="", layout="wide")


# ══════════════════════════════════════════════════════════════════════════════
# OLLAMA HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def ollama_call(prompt: str, timeout: int = 60) -> tuple:
    model   = _get_model()
    payload = json.dumps({
        "model": model, "prompt": prompt, "stream": False,
        "options": {"temperature": 0.4, "num_predict": 400, "top_p": 0.9},
    }).encode()
    req = urllib.request.Request(
        f"{OLLAMA_BASE}/api/generate", data=payload,
        headers={"Content-Type": "application/json"}, method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout) as r:
            result = json.loads(r.read().decode()).get("response", "").strip()
            if not result:
                return None, "Model returned empty response"
            return result, None
    except urllib.error.URLError as e:
        return None, f"Ollama not running — run: ollama serve  ({e})"
    except Exception as e:
        return None, f"Ollama error: {e}"


def ollama_json(prompt: str, max_tokens: int = 600) -> tuple:
    model   = _get_model()
    payload = json.dumps({
        "model": model, "prompt": prompt, "stream": False,
        "options": {"temperature": 0.1, "num_predict": max_tokens, "top_p": 0.9},
    }).encode()
    req = urllib.request.Request(
        f"{OLLAMA_BASE}/api/generate", data=payload,
        headers={"Content-Type": "application/json"}, method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=90) as r:
            text = json.loads(r.read().decode()).get("response", "").strip()
    except Exception as e:
        return {}, f"Ollama error: {e}"
    if not text:
        return {}, "Empty response from model"
    clean = re.sub(r"```(?:json)?", "", text, flags=re.IGNORECASE).strip().strip("`").strip()
    first = clean.find("{")
    last  = clean.rfind("}")
    if first != -1 and last != -1 and last > first:
        clean = clean[first:last + 1]
    else:
        return {}, f"No JSON found. Raw: {text[:300]}"
    try:
        return json.loads(clean), None
    except json.JSONDecodeError:
        repaired      = re.sub(r",\s*$", "", clean.rstrip())
        open_braces   = repaired.count("{") - repaired.count("}")
        open_brackets = repaired.count("[") - repaired.count("]")
        for _ in range(open_brackets): repaired += "]"
        for _ in range(open_braces):   repaired += "}"
        try:
            return json.loads(repaired), None
        except Exception as pe2:
            return {}, f"JSON parse failed: {pe2} | raw: {text[:200]}"


# ══════════════════════════════════════════════════════════════════════════════
# FILE READERS
# ══════════════════════════════════════════════════════════════════════════════

def read_pdf(f) -> str:
    try:
        reader = PdfReader(f)
        return "\n".join(p.extract_text() for p in reader.pages if p.extract_text()).strip()
    except Exception:
        return ""

def read_docx_text(f) -> str:
    try:
        doc = Document(BytesIO(f) if isinstance(f, bytes) else BytesIO(f.read()))
        return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""

def read_resume_file(uploaded) -> str:
    name = uploaded.name.lower()
    if name.endswith(".pdf"):  return read_pdf(uploaded)
    if name.endswith(".docx"): return read_docx_text(uploaded)
    raw = uploaded.read()
    try:    return raw.decode("utf-8").strip()
    except: return raw.decode("latin-1").strip()


# ══════════════════════════════════════════════════════════════════════════════
# DOCX WRITE-BACK
# ══════════════════════════════════════════════════════════════════════════════

def _all_paras(document):
    result = []
    for child in document.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            result.append(Paragraph(child, document))
        elif tag == "tbl":
            tbl = Table(child, document)
            for row in tbl.rows:
                for cell in row.cells:
                    result.extend(cell.paragraphs)
    return result

def _set_para_text(para, new_text: str):
    from docx.oxml.ns import qn
    p_elem   = para._p
    pPr      = p_elem.find(qn("w:pPr"))
    pPr_copy = copy.deepcopy(pPr) if pPr is not None else None
    first_r  = p_elem.find(qn("w:r"))
    rPr_copy = None
    if first_r is not None:
        rPr = first_r.find(qn("w:rPr"))
        if rPr is not None:
            rPr_copy = copy.deepcopy(rPr)
    for child in list(p_elem): p_elem.remove(child)
    if pPr_copy is not None: p_elem.append(pPr_copy)
    from lxml import etree
    r_elem = etree.SubElement(p_elem, qn("w:r"))
    if rPr_copy is not None: r_elem.append(rPr_copy)
    t_elem = etree.SubElement(r_elem, qn("w:t"))
    t_elem.text = new_text
    if new_text.startswith(" ") or new_text.endswith(" "):
        t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

def inject_text_into_docx(docx_bytes: bytes, new_text: str) -> bytes:
    doc       = Document(BytesIO(docx_bytes))
    paras     = [p for p in _all_paras(doc) if p.text.strip()]
    new_lines = [ln for ln in new_text.splitlines() if ln.strip()]
    for i, para in enumerate(paras):
        if i < len(new_lines):
            _set_para_text(para, new_lines[i])
    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════════════════════════

def init_state():
    defaults = {
        "rb_name": "", "rb_email": "", "rb_phone": "",
        "rb_location": "", "rb_linkedin": "", "rb_github": "", "rb_portfolio": "",
        "rb_summary": "", "rb_skills": "",
        "rb_experience": [], "rb_projects": [], "rb_education": [], "rb_certs": [],
        "rb_template": "Modern",
        "rb_jd_company": "", "rb_jd_role": "", "rb_jd_text": "",
        "rb_analysis": {}, "rb_imported": False, "rb_ats_result": {},
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_state()


# ══════════════════════════════════════════════════════════════════════════════
# AI FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════

def ai_suggest_summary(name, role, skills, jd_text):
    role  = role or "Software Engineering Intern"
    skls  = (skills or "Python, problem solving")[:100]
    jd_kw = " ".join((jd_text or "")[:200].split()[:25])
    p = (
        f"Write a 3-sentence professional resume summary.\n"
        f"Person: {name or 'Student'}, Role: {role}, Skills: {skls}\n"
        f"Use JD keywords: {jd_kw}\n"
        f"Format: Sentence1=who+role. Sentence2=skills+achievement. Sentence3=goal.\n"
        f"RETURN ONLY the paragraph. No labels. No quotes. No preamble."
    )
    t, e = ollama_call(p, 80)
    return (t or "").strip(), e

def ai_suggest_bullets(role_title, company, skills_used, jd_text, existing=""):
    tech  = (skills_used or "Python")[:80]
    jd_kw = " ".join((jd_text or "")[:150].split()[:20])
    p = (
        f"Write 3 resume bullet points for {role_title or 'Intern'} at {company or 'Company'}.\n"
        f"Tech: {tech}. JD keywords: {jd_kw}\n"
        f"IMPORTANT: Output ONLY these 3 lines:\n"
        f"- [Past-tense verb] [what] using [tech] [result]\n"
        f"- [Past-tense verb] [what] using [tech] [result]\n"
        f"- [Past-tense verb] [what] using [tech] [result]"
    )
    t, e = ollama_call(p, 100)
    return (t or "").strip(), e

def ai_suggest_project_bullets(project_name, tech_stack, jd_text):
    tech = (tech_stack or "Python")[:80]
    p = (
        f"Write 3 resume bullets for project: {project_name}. Tech: {tech}\n"
        f"ONLY output these 3 lines:\n"
        f"- [verb] [feature] using [tech] [outcome]\n"
        f"- [verb] [feature] using [tech] [outcome]\n"
        f"- [verb] [feature] using [tech] [outcome]"
    )
    t, e = ollama_call(p, 100)
    return (t or "").strip(), e

def ai_suggest_skills(jd_text, current_skills):
    jd_kw = " ".join((jd_text or "")[:300].split()[:40])
    cur   = (current_skills or "")[:80]
    p = (
        f"List 8 technical skills for: {jd_kw}\n"
        f"Exclude: {cur}\n"
        f"ONLY output: Skill1, Skill2, Skill3, Skill4, Skill5, Skill6, Skill7, Skill8"
    )
    t, e = ollama_call(p, 60)
    return (t or "").strip(), e

def ai_full_analysis(resume_text, role, jd_text):
    res = " ".join((resume_text or "").split()[:300])
    jd  = " ".join((jd_text    or "").split()[:150])
    p = (
        f"Analyse resume vs job. ONLY output valid JSON:\n"
        f'{{"match_score":70,"ats_score":65,'
        f'"missing_keywords":["k1","k2","k3"],'
        f'"matched_keywords":["m1","m2"],'
        f'"strengths":["s1","s2"],'
        f'"critical_gaps":["g1","g2"],'
        f'"quick_wins":["t1","t2"],'
        f'"section_scores":{{"summary":7,"skills":8,"experience":7,"projects":6,"education":9}},'
        f'"interview_questions":["q1","q2","q3"],'
        f'"linkedin_headline":"headline"}}\n'
        f"RESUME: {res}\nROLE: {role or 'Intern'}\nJD: {jd}"
    )
    return ollama_json(p, 600)

# ── NEW: Skill Roadmap ────────────────────────────────────────────────────────
def ai_skill_roadmap(target_role, weeks, current_skills, jd_text):
    cur = (current_skills or "beginner")[:300]
    jd  = " ".join((jd_text or "")[:300].split()[:40])
    prompt = (
        f"Create a {weeks}-week skill development roadmap.\n"
        f"Target: {target_role}\n"
        f"Current skills: {cur}\n"
        f"JD keywords: {jd if jd else 'not provided'}\n\n"
        f"Format as week-by-week plan. Each week must follow exactly this structure:\n"
        f"Week N: [Theme]\n"
        f"- Action item 1 (resource: specific course/book)\n"
        f"- Action item 2\n"
        f"- Project: [mini project to build]\n\n"
        f"Focus on skills with highest ROI for internship applications.\n"
        f"Keep it practical, concise, and tailored for Indian tech internships.\n"
        f"RETURN ONLY THE ROADMAP."
    )
    text, err = ollama_call(prompt, 120)
    return (text or "").strip(), err


# ══════════════════════════════════════════════════════════════════════════════
# REGEX RESUME PARSER
# ══════════════════════════════════════════════════════════════════════════════

SECTION_NAMES = {
    "experience":     ["experience","work experience","employment","internship","internships","career"],
    "education":      ["education","academic","qualifications","schooling"],
    "skills":         ["skills","technical skills","technologies","competencies","expertise"],
    "projects":       ["projects","personal projects","academic projects","portfolio"],
    "certifications": ["certifications","certificates","courses","training","achievements"],
    "summary":        ["summary","objective","profile","about","professional summary"],
}

def _regex_parse_resume(text: str) -> dict:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    full  = "\n".join(lines)

    email = phone = linkedin = github = name = location = ""

    m = re.search(r"[\w.+-]+@[\w-]+\.[a-z]{2,6}", full, re.I)
    if m: email = m.group()

    m = re.search(r"(\+?[\d][\d\s\-().]{8,15}\d)", full)
    if m: phone = re.sub(r"\s+", " ", m.group().strip())

    m = re.search(r"linkedin\.com/in/([\w\-]+)", full, re.I)
    if m: linkedin = "linkedin.com/in/" + m.group(1)

    m = re.search(r"github\.com/([\w\-]+)", full, re.I)
    if m: github = "github.com/" + m.group(1)

    for ln in lines[:6]:
        if re.search(r"@|http|linkedin|github|phone|email|mobile|\d{6}", ln, re.I):
            continue
        if re.match(r"^[A-Z][a-z]+([ ][A-Z][a-z]+){0,3}$", ln):
            name = ln
            break

    for ln in lines[:10]:
        m = re.search(r"\b(Bengaluru|Bangalore|Mumbai|Delhi|Hyderabad|Pune|Chennai|India|Remote)\b", ln, re.I)
        if m:
            location = ln[:60]
            break

    sec_lines = {k: [] for k in SECTION_NAMES}
    current   = None
    for ln in lines:
        low   = ln.lower().rstrip(":").strip()
        found = False
        for sec, kws in SECTION_NAMES.items():
            if low in kws:
                current = sec
                found   = True
                break
        if not found and current:
            sec_lines[current].append(ln)

    summary     = " ".join(sec_lines.get("summary", []))[:600]
    raw_skills  = " ".join(sec_lines.get("skills", []))
    skill_items = [s.strip().rstrip(",.") for s in re.split(r"[,|•\n]", raw_skills) if s.strip() and len(s.strip()) > 1]
    skills_str  = ", ".join(skill_items[:30])

    exp_list = []
    exp_text = "\n".join(sec_lines.get("experience", []))
    for block in re.split(r"\n{2,}", exp_text):
        blines = [l.strip() for l in block.splitlines() if l.strip()]
        if not blines: continue
        title, company, start, end, loc, bullets = "", "", "", "", "", []
        for bl in blines:
            if not title and re.search(r"intern|engineer|developer|analyst|manager|lead|consultant|scientist", bl, re.I):
                title = bl[:80]
            elif not company and re.match(r"[A-Z]", bl) and len(bl) < 60:
                company = bl[:60]
            elif re.search(r"\d{4}", bl):
                dates = re.findall(r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Present|\d{4})", bl)
                if len(dates) >= 2: start, end = dates[0], dates[-1]
                elif len(dates) == 1: start = dates[0]
            elif bl.startswith(("•", "-", "*")):
                bullets.append(bl.lstrip("•-* "))
        if title or company:
            exp_list.append({"title": title, "company": company, "start": start, "end": end, "location": loc, "bullets": bullets[:4] or ["","",""]})

    proj_list = []
    proj_text = "\n".join(sec_lines.get("projects", []))
    for block in re.split(r"\n{2,}", proj_text):
        blines = [l.strip() for l in block.splitlines() if l.strip()]
        if not blines: continue
        name_p, tech, link, bullets = blines[0][:80] if blines else "", "", "", []
        for bl in blines[1:]:
            m = re.search(r"(https?://\S+|github\.com/\S+)", bl, re.I)
            if m: link = m.group(1)
            elif re.search(r"python|react|node|java|sql|mongodb|docker|flask|django|ml|ai", bl, re.I):
                tech = bl[:80]
            elif bl.startswith(("•", "-", "*")):
                bullets.append(bl.lstrip("•-* "))
        if name_p:
            proj_list.append({"name": name_p, "tech": tech, "link": link, "bullets": bullets[:3] or ["",""]})

    edu_list = []
    edu_text = "\n".join(sec_lines.get("education", []))
    for block in re.split(r"\n{2,}", edu_text):
        blines = [l.strip() for l in block.splitlines() if l.strip()]
        if not blines: continue
        degree, institution, year, gpa = "", "", "", ""
        for bl in blines:
            if re.search(r"b\.?tech|m\.?tech|b\.?e|bachelor|master|bca|mca|bsc|msc|phd|diploma", bl, re.I):
                degree = bl[:80]
            elif re.search(r"university|institute|college|iit|nit|bits", bl, re.I):
                institution = bl[:80]
            elif re.search(r"\d{4}", bl):
                year = bl[:20]
            elif re.search(r"\bgpa\b|\bcgpa\b|%|grade", bl, re.I):
                gpa = bl[:20]
        if degree or institution:
            edu_list.append({"degree": degree, "institution": institution, "year": year, "gpa": gpa})

    cert_list = [{"name": cl[:80], "issuer": "", "year": ""} for cl in sec_lines.get("certifications", []) if len(cl) > 3]

    return {
        "name": name, "email": email, "phone": phone,
        "location": location, "linkedin": linkedin, "github": github, "portfolio": "",
        "summary": summary, "skills": skills_str,
        "experience": exp_list, "projects": proj_list,
        "education": edu_list, "certs": cert_list,
    }


# ══════════════════════════════════════════════════════════════════════════════
# RESUME TEXT BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_resume_text() -> str:
    s     = st.session_state
    lines = []

    if s["rb_name"]: lines.append(s["rb_name"])
    contact = " | ".join(x for x in [s["rb_email"], s["rb_phone"], s["rb_location"]] if x)
    if contact: lines.append(contact)
    links = " | ".join(x for x in [s["rb_linkedin"], s["rb_github"], s["rb_portfolio"]] if x)
    if links: lines.append(links)
    lines.append("")

    if s["rb_summary"]:
        lines += ["PROFESSIONAL SUMMARY", "-" * 40, s["rb_summary"], ""]

    if s["rb_skills"]:
        lines += ["SKILLS", "-" * 40, s["rb_skills"], ""]

    if s["rb_experience"]:
        lines += ["WORK EXPERIENCE", "-" * 40]
        for exp in s["rb_experience"]:
            header = " | ".join(x for x in [exp.get("title",""), exp.get("company",""), exp.get("start",""), exp.get("end",""), exp.get("location","")] if x)
            if header: lines.append(header)
            for b in exp.get("bullets", []):
                if b.strip(): lines.append(f"  - {b.strip().lstrip('-').strip()}")
        lines.append("")

    if s["rb_projects"]:
        lines += ["PROJECTS", "-" * 40]
        for proj in s["rb_projects"]:
            header = " | ".join(x for x in [proj.get("name",""), proj.get("tech",""), proj.get("link","")] if x)
            if header: lines.append(header)
            for b in proj.get("bullets", []):
                if b.strip(): lines.append(f"  - {b.strip().lstrip('-').strip()}")
        lines.append("")

    if s["rb_education"]:
        lines += ["EDUCATION", "-" * 40]
        for edu in s["rb_education"]:
            row = " | ".join(x for x in [edu.get("degree",""), edu.get("institution",""), edu.get("year",""), edu.get("gpa","")] if x)
            if row: lines.append(row)
        lines.append("")

    if s["rb_certs"]:
        lines += ["CERTIFICATIONS", "-" * 40]
        for c in s["rb_certs"]:
            row = " | ".join(x for x in [c.get("name",""), c.get("issuer",""), c.get("year","")] if x)
            if row: lines.append(row)
        lines.append("")

    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# HTML TEMPLATES
# ══════════════════════════════════════════════════════════════════════════════

def _esc(s): return str(s).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

def render_html_template(template: str) -> str:
    s         = st.session_state
    name      = _esc(s["rb_name"] or "Your Name")
    email     = _esc(s["rb_email"])
    phone     = _esc(s["rb_phone"])
    location  = _esc(s["rb_location"])
    linkedin  = _esc(s["rb_linkedin"])
    github    = _esc(s["rb_github"])
    portfolio = _esc(s["rb_portfolio"])
    summary   = _esc(s["rb_summary"])
    skills    = _esc(s["rb_skills"])

    contact_html = " &nbsp;|&nbsp; ".join(x for x in [email, phone, location] if x)
    links_html   = ""
    if linkedin:  links_html += f'<a href="https://{linkedin}" target="_blank">LinkedIn</a> '
    if github:    links_html += f'<a href="https://{github}"   target="_blank">GitHub</a> '
    if portfolio: links_html += f'<a href="{portfolio}"        target="_blank">Portfolio</a> '

    skill_pills = ""
    if skills:
        for sk in re.split(r"[,|]", skills):
            sk = sk.strip()
            if sk: skill_pills += f'<span class="skill-pill">{_esc(sk)}</span>'

    exp_html = ""
    for exp in s["rb_experience"]:
        bullets_html = "".join(
            f"<li>{_esc(b.strip().lstrip('-').strip())}</li>"
            for b in exp.get("bullets", []) if b.strip()
        )
        dates = " &nbsp;·&nbsp; ".join(filter(None, [_esc(exp.get("start","")), _esc(exp.get("end","")), _esc(exp.get("location",""))]))
        exp_html += (
            f'<div class="entry"><div class="entry-header"><div>'
            f'<span class="entry-title">{_esc(exp.get("title",""))}</span>'
            f' &nbsp;@&nbsp; <span class="entry-sub">{_esc(exp.get("company",""))}</span></div>'
            f'<span class="entry-date">{dates}</span></div>'
            f'{"<ul class=bullets>"+bullets_html+"</ul>" if bullets_html else ""}</div>'
        )

    proj_html = ""
    for proj in s["rb_projects"]:
        bullets_html = "".join(
            f"<li>{_esc(b.strip().lstrip('-').strip())}</li>"
            for b in proj.get("bullets", []) if b.strip()
        )
        tech      = proj.get("tech", "")
        link      = proj.get("link", "")
        link_html = f'<a href="{_esc(link)}" target="_blank" class="proj-link">{_esc(link)}</a>' if link else ""
        proj_html += (
            f'<div class="entry"><div class="entry-header"><div>'
            f'<span class="entry-title">{_esc(proj.get("name",""))}</span>'
            f'{f"<span class=tech-stack>{_esc(tech)}</span>" if tech else ""}</div>'
            f'{link_html}</div>'
            f'{"<ul class=bullets>"+bullets_html+"</ul>" if bullets_html else ""}</div>'
        )

    edu_html = ""
    for edu in s["rb_education"]:
        meta = " &nbsp;·&nbsp; ".join(filter(None, [_esc(edu.get("year","")), _esc(edu.get("gpa",""))]))
        edu_html += (
            f'<div class="entry"><div class="entry-header">'
            f'<span class="entry-title">{_esc(edu.get("degree",""))}</span>'
            f'<span class="entry-date">{meta}</span></div>'
            f'<div class="entry-sub">{_esc(edu.get("institution",""))}</div></div>'
        )

    cert_html = "".join(
        f'<div class="cert-item"><b>{_esc(c.get("name",""))}</b>'
        f' — {_esc(c.get("issuer",""))} <span class="entry-date">{_esc(c.get("year",""))}</span></div>'
        for c in s["rb_certs"]
    )

    if template == "Modern":
        accent="#2563eb"; header_bg="#1e293b"; header_text="#ffffff"; body_bg="#ffffff"
        section_color="#2563eb"; font="'Inter','Segoe UI',sans-serif"
        header_style =f"background:{header_bg};color:{header_text};padding:32px 40px 28px;"
        name_style   ="font-size:2rem;font-weight:800;letter-spacing:1px;"
        sec_border   =f"border-left:4px solid {accent};padding-left:12px;"
    elif template == "Professional":
        accent="#1a4e8a"; header_bg="#ffffff"; header_text="#1a2744"; body_bg="#ffffff"
        section_color="#1a4e8a"; font="'Georgia','Times New Roman',serif"
        header_style =f"border-bottom:3px solid {accent};padding:24px 40px 20px;background:#ffffff;"
        name_style   =f"font-size:1.9rem;font-weight:700;color:{header_text};letter-spacing:0.5px;"
        sec_border   =f"border-bottom:2px solid {accent};padding-bottom:4px;margin-bottom:10px;"
    else:
        accent="#111827"; header_bg="#f9fafb"; header_text="#111827"; body_bg="#ffffff"
        section_color="#374151"; font="'IBM Plex Sans','Helvetica Neue',sans-serif"
        header_style =f"background:{header_bg};padding:28px 40px 20px;border-bottom:1px solid #e5e7eb;"
        name_style   ="font-size:1.8rem;font-weight:600;letter-spacing:0.5px;"
        sec_border   ="border-bottom:1px solid #d1d5db;padding-bottom:3px;margin-bottom:10px;"

    return f"""<!DOCTYPE html><html lang="en"><head><meta charset="utf-8">
<title>{name} — Resume</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0;}}
body{{font-family:{font};background:#e5e7eb;color:#1f2937;font-size:14px;line-height:1.55;}}
.page{{max-width:800px;margin:0 auto;background:{body_bg};box-shadow:0 4px 24px rgba(0,0,0,.15);min-height:1100px;}}
.resume-header{{ {header_style} }}
.resume-name{{ {name_style} color:{header_text}; }}
.resume-contact{{margin-top:6px;font-size:0.82rem;opacity:0.85;color:{header_text};}}
.resume-links{{margin-top:5px;font-size:0.8rem;}}
.resume-links a{{color:{accent};text-decoration:none;margin-right:14px;}}
.resume-body{{padding:28px 40px;}}
.section{{margin-bottom:22px;}}
.section-title{{font-size:0.82rem;font-weight:800;text-transform:uppercase;letter-spacing:1.2px;color:{section_color};margin-bottom:10px;{sec_border}}}
.entry{{margin-bottom:12px;}}
.entry-header{{display:flex;justify-content:space-between;align-items:baseline;flex-wrap:wrap;gap:4px;}}
.entry-title{{font-weight:700;font-size:0.93rem;color:#111827;}}
.entry-sub{{font-size:0.84rem;color:#4b5563;margin-top:1px;}}
.entry-date{{font-size:0.77rem;color:#6b7280;white-space:nowrap;}}
.tech-stack{{font-size:0.73rem;background:#f3f4f6;border:1px solid #e5e7eb;border-radius:4px;padding:1px 7px;margin-left:8px;color:#4b5563;}}
.proj-link{{font-size:0.72rem;color:{accent};text-decoration:none;}}
.bullets{{margin:6px 0 0 16px;}}
.bullets li{{font-size:0.84rem;color:#374151;margin-bottom:3px;}}
.skill-pill{{display:inline-block;background:#f0f4ff;border:1px solid #c7d7f7;border-radius:20px;padding:3px 12px;margin:3px 4px 3px 0;font-size:0.78rem;font-weight:600;color:{accent};}}
.cert-item{{font-size:0.83rem;color:#374151;margin-bottom:5px;}}
.summary-text{{font-size:0.86rem;color:#374151;line-height:1.65;}}
@media print{{body{{background:white;}}.page{{box-shadow:none;max-width:100%;}}}}
</style></head><body><div class="page">
  <div class="resume-header">
    <div class="resume-name">{name}</div>
    <div class="resume-contact">{contact_html}</div>
    {"<div class='resume-links'>"+links_html+"</div>" if links_html else ""}
  </div>
  <div class="resume-body">
    {"<div class='section'><div class='section-title'>Professional Summary</div><p class='summary-text'>"+summary+"</p></div>" if summary else ""}
    {"<div class='section'><div class='section-title'>Skills</div>"+skill_pills+"</div>" if skill_pills else ""}
    {"<div class='section'><div class='section-title'>Work Experience</div>"+exp_html+"</div>" if exp_html else ""}
    {"<div class='section'><div class='section-title'>Projects</div>"+proj_html+"</div>" if proj_html else ""}
    {"<div class='section'><div class='section-title'>Education</div>"+edu_html+"</div>" if edu_html else ""}
    {"<div class='section'><div class='section-title'>Certifications</div>"+cert_html+"</div>" if cert_html else ""}
  </div>
</div></body></html>"""


# ══════════════════════════════════════════════════════════════════════════════
# IMPORT / PARSE
# ══════════════════════════════════════════════════════════════════════════════

def import_resume_via_ai(resume_text: str):
    if not resume_text or len(resume_text.strip()) < 30:
        return "Could not extract text — PDF may be scanned/image-based."
    data = _regex_parse_resume(resume_text)
    s = st.session_state
    s["rb_name"]       = data.get("name", "")
    s["rb_email"]      = data.get("email", "")
    s["rb_phone"]      = data.get("phone", "")
    s["rb_location"]   = data.get("location", "")
    s["rb_linkedin"]   = data.get("linkedin", "")
    s["rb_github"]     = data.get("github", "")
    s["rb_portfolio"]  = data.get("portfolio", "")
    s["rb_summary"]    = data.get("summary", "")
    s["rb_skills"]     = data.get("skills", "")
    s["rb_experience"] = data.get("experience", [])
    s["rb_projects"]   = data.get("projects", [])
    s["rb_education"]  = data.get("education", [])
    s["rb_certs"]      = data.get("certs", [])
    s["rb_imported"]   = True
    return None


# ══════════════════════════════════════════════════════════════════════════════
# UI HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def section_header(title: str):
    st.markdown(
        f"<div style='margin:18px 0 8px;'>"
        f"<span style='font-size:1rem;font-weight:800;color:#e8eef8;'>{title}</span>"
        f"</div>",
        unsafe_allow_html=True,
    )

def pill_badge(text, color="#4a90d9"):
    return (
        f"<span style='display:inline-block;background:{color}18;border:1px solid {color}55;"
        f"border-radius:20px;padding:2px 10px;font-size:0.72rem;font-weight:700;"
        f"color:{color};margin:2px 3px 2px 0;'>{text}</span>"
    )

def score_color(v):
    if v >= 75: return "#22c55e"
    if v >= 50: return "#f59e0b"
    return "#ef4444"

def score_label(v):
    if v >= 75: return "Strong"
    if v >= 50: return "Good"
    return "Needs Work"

def _build_parsed_data() -> dict:
    return {
        "name":       st.session_state.get("rb_name", ""),
        "summary":    st.session_state.get("rb_summary", ""),
        "skills":     st.session_state.get("rb_skills", ""),
        "experience": st.session_state.get("rb_experience", []),
        "projects":   st.session_state.get("rb_projects", []),
        "education":  st.session_state.get("rb_education", []),
        "certs":      st.session_state.get("rb_certs", []),
    }

def _run_ats_score():
    if _ATS_AVAILABLE:
        st.session_state["rb_ats_result"] = score_resume(
            build_resume_text(), _build_parsed_data()
        )


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    _PENDING = {
        "_pending_rb_summary": "rb_summary", "_pending_rb_skills":   "rb_skills",
        "_pending_rb_name":    "rb_name",    "_pending_rb_email":    "rb_email",
        "_pending_rb_phone":   "rb_phone",   "_pending_rb_location": "rb_location",
        "_pending_rb_linkedin":"rb_linkedin", "_pending_rb_github":  "rb_github",
        "_pending_rb_portfolio":"rb_portfolio",
    }
    flushed = False
    for pk, wk in _PENDING.items():
        if pk in st.session_state:
            st.session_state[wk] = st.session_state.pop(pk)
            flushed = True

    if flushed and st.session_state.get("rb_ats_result"):
        _run_ats_score()

    if st.session_state.pop("_pending_rescore", False):
        _run_ats_score()

    st.markdown("""<style>
    [data-testid="stSidebar"] { background: rgba(5,12,30,0.98); }
    .block-container { padding-top: 1.5rem; }
    div[data-testid="stTextInput"] input,
    div[data-testid="stTextArea"] textarea {
        background: rgba(10,20,50,0.7) !important;
        border: 1px solid rgba(74,144,217,0.3) !important;
        color: #e8eef8 !important; border-radius: 8px !important;
    }
    div[data-testid="stTextInput"] input:focus,
    div[data-testid="stTextArea"] textarea:focus {
        border-color: rgba(74,144,217,0.8) !important;
        box-shadow: 0 0 0 2px rgba(74,144,217,0.15) !important;
    }
    .stButton > button { border-radius: 8px !important; font-weight: 700 !important; }
    .stTabs [data-baseweb="tab"] { font-weight: 700 !important; font-size: 0.9rem !important; }
    </style>""", unsafe_allow_html=True)

    with st.sidebar:
        st.markdown(
            "<div style='font-size:1rem;font-weight:800;color:#e8eef8;margin-bottom:12px;'>"
            "Ollama</div>",
            unsafe_allow_html=True,
        )
        ok, models, _ = _ollama_available()
        if ok and models:
            st.success("Running")
            st.selectbox("Model", options=models, index=0, key="ollama_model_pick")
        else:
            st.error("Not detected")
            st.markdown(
                "<div style='font-size:0.74rem;color:#f87171;line-height:1.8;'>"
                "Start Ollama, then run:<br>"
                "<code style='background:#1a1a2e;padding:2px 6px;border-radius:3px;'>"
                "ollama pull phi3:mini</code></div>",
                unsafe_allow_html=True,
            )
            st.text_input("Model name", value="phi3:mini", key="ollama_model_manual")

    st.markdown(
        "<h1 style='margin:0 0 4px;font-size:1.8rem;font-weight:900;color:#e8eef8;'>"
        "Resume Builder</h1>"
        "<p style='color:#7a9ec0;font-size:0.85rem;margin:0 0 16px;'>"
        "Build or import your resume — AI fills sections and analyses JD match via Ollama.</p>",
        unsafe_allow_html=True,
    )

    imp_col, tpl_col, _ = st.columns([2, 2, 3])
    with imp_col:
        up = st.file_uploader("Import resume", type=["pdf","docx","txt"],
                              key="import_file", label_visibility="collapsed")
        if up:
            if st.button("Parse & Score Resume", key="parse_btn", use_container_width=True, type="primary"):
                with st.spinner("Reading…"):
                    raw = read_resume_file(up)
                if not raw or len(raw.strip()) < 30:
                    st.error("Could not extract text — PDF may be scanned.\n\n**Fix:** Copy text from PDF and use Paste Text tab.")
                else:
                    with st.spinner(f"Parsing {len(raw.split())} words…"):
                        err = import_resume_via_ai(raw)
                    if err:
                        st.error(f"Parse error: {err}")
                    else:
                        _run_ats_score()
                        st.rerun()
    with tpl_col:
        st.session_state["rb_template"] = st.selectbox(
            "Template", ["Modern","Professional","Minimal"],
            index=["Modern","Professional","Minimal"].index(st.session_state["rb_template"]),
            key="tpl_select", label_visibility="collapsed",
        )

    ats_result = st.session_state.get("rb_ats_result", {})
    if ats_result and _ATS_AVAILABLE:
        render_ats_card(ats_result)
        if st.button("Re-score", key="rescore_btn"):
            _run_ats_score()
            st.rerun()

    st.markdown("---")

    # ── Tabs — 4 total ────────────────────────────────────────────────────
    tab_build, tab_ai, tab_preview, tab_roadmap = st.tabs(
        ["Build Resume", "AI Analyse & Improve", "Preview & Export", "Skill Roadmap"]
    )

    # ══════════════════════════════════════════════════════════════════════
    # TAB 1 — BUILD
    # ══════════════════════════════════════════════════════════════════════
    with tab_build:
        with st.expander("Target Internship (for AI suggestions)", expanded=False):
            jc1, jc2 = st.columns(2)
            with jc1: st.text_input("Role",    placeholder="e.g. Software Engineering Intern", key="rb_jd_role")
            with jc2: st.text_input("Company", placeholder="e.g. Google India",                key="rb_jd_company")
            st.text_area("Job Description", height=100, key="rb_jd_text",
                         placeholder="Paste the full JD for best AI suggestions")

        section_header("Personal Information")
        p1, p2, p3 = st.columns(3)
        with p1: st.text_input("Full Name *", key="rb_name",  placeholder="Arjun Sharma")
        with p2: st.text_input("Email *",     key="rb_email", placeholder="arjun@email.com")
        with p3: st.text_input("Phone",       key="rb_phone", placeholder="+91 98765 43210")
        p4, p5, p6, p7 = st.columns(4)
        with p4: st.text_input("Location",  key="rb_location",  placeholder="Bengaluru, India")
        with p5: st.text_input("LinkedIn",  key="rb_linkedin",  placeholder="linkedin.com/in/arjun")
        with p6: st.text_input("GitHub",    key="rb_github",    placeholder="github.com/arjun")
        with p7: st.text_input("Portfolio", key="rb_portfolio", placeholder="arjun.dev")

        section_header("Professional Summary")
        sc, bc = st.columns([5, 1])
        with sc:
            st.text_area("Summary", height=100, key="rb_summary",
                         placeholder="3-sentence pitch tailored to your target role…",
                         label_visibility="collapsed")
        with bc:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("AI Write", key="ai_summary", use_container_width=True):
                with st.spinner("Writing…"):
                    text, err = ai_suggest_summary(st.session_state["rb_name"], st.session_state["rb_jd_role"],
                                                   st.session_state["rb_skills"], st.session_state["rb_jd_text"])
                if err: st.error(err)
                else:
                    st.session_state["_pending_rb_summary"] = text
                    st.session_state["_pending_rescore"] = True
                    st.rerun()

        section_header("Skills")
        sk, skb = st.columns([5, 1])
        with sk:
            st.text_area("Skills", height=80, key="rb_skills",
                         placeholder="Python, React, Node.js, SQL, Git, Docker…",
                         label_visibility="collapsed")
        with skb:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("AI Suggest", key="ai_skills", use_container_width=True):
                if not st.session_state["rb_jd_text"]:
                    st.warning("Paste JD above first.")
                else:
                    with st.spinner("Suggesting…"):
                        text, err = ai_suggest_skills(st.session_state["rb_jd_text"], st.session_state["rb_skills"])
                    if err: st.error(err)
                    else:
                        existing = st.session_state["rb_skills"].rstrip(", ")
                        st.session_state["_pending_rb_skills"] = (existing + ", " + text).strip(", ")
                        st.session_state["_pending_rescore"] = True
                        st.rerun()

        section_header("Work Experience")
        if st.button("+ Add Experience", key="add_exp"):
            st.session_state["rb_experience"].append({"title":"","company":"","start":"","end":"","location":"","bullets":["","",""]})
            st.rerun()

        for i, exp in enumerate(st.session_state["rb_experience"]):
            lbl = f"{exp['title']} @ {exp['company']}" if exp["title"] else f"Experience {i+1}"
            with st.expander(lbl, expanded=True):
                e1, e2, e3, e4, e5 = st.columns([2,2,1,1,1])
                with e1: exp["title"]    = st.text_input("Job Title", value=exp.get("title",""),    key=f"e_title_{i}",   placeholder="SWE Intern")
                with e2: exp["company"]  = st.text_input("Company",   value=exp.get("company",""),  key=f"e_company_{i}", placeholder="Google India")
                with e3: exp["start"]    = st.text_input("Start",     value=exp.get("start",""),    key=f"e_start_{i}",   placeholder="Jun 2024")
                with e4: exp["end"]      = st.text_input("End",       value=exp.get("end",""),      key=f"e_end_{i}",     placeholder="Aug 2024")
                with e5: exp["location"] = st.text_input("Location",  value=exp.get("location",""), key=f"e_loc_{i}",     placeholder="Bangalore")
                st.caption("Bullet points")
                bullets = exp.get("bullets", ["","",""])
                while len(bullets) < 3: bullets.append("")
                new_bullets = []
                for j, b in enumerate(bullets):
                    new_bullets.append(st.text_input(f"Bullet {j+1}", value=b, key=f"e_bullet_{i}_{j}",
                                                     placeholder="Action verb + what + result", label_visibility="collapsed"))
                exp["bullets"] = new_bullets
                bc1, bc2, bc3 = st.columns([2,2,1])
                with bc1: tech_used = st.text_input("Tech used", key=f"e_tech_{i}", placeholder="Python, Django, PostgreSQL")
                with bc2:
                    if st.button("AI Write Bullets", key=f"ai_exp_{i}", use_container_width=True):
                        with st.spinner("Writing…"):
                            text, err = ai_suggest_bullets(exp.get("title","Intern"), exp.get("company","Company"),
                                                           tech_used or "", st.session_state["rb_jd_text"],
                                                           "\n".join(b for b in new_bullets if b))
                        if err: st.error(err)
                        else:
                            parsed = [ln.strip().lstrip("-").strip() for ln in text.splitlines() if ln.strip()][:3]
                            while len(parsed) < 3: parsed.append("")
                            exp["bullets"] = parsed
                            st.session_state["_pending_rescore"] = True
                            st.rerun()
                with bc3:
                    if st.button("Remove", key=f"del_exp_{i}", use_container_width=True):
                        st.session_state["rb_experience"].pop(i)
                        st.rerun()

        section_header("Projects")
        if st.button("+ Add Project", key="add_proj"):
            st.session_state["rb_projects"].append({"name":"","tech":"","link":"","bullets":["",""]})
            st.rerun()

        for i, proj in enumerate(st.session_state["rb_projects"]):
            lbl = proj["name"] if proj["name"] else f"Project {i+1}"
            with st.expander(lbl, expanded=True):
                pr1, pr2, pr3 = st.columns([2,2,2])
                with pr1: proj["name"] = st.text_input("Project Name", value=proj.get("name",""), key=f"p_name_{i}", placeholder="AI Chatbot")
                with pr2: proj["tech"] = st.text_input("Tech Stack",   value=proj.get("tech",""), key=f"p_tech_{i}", placeholder="Python, FastAPI, React")
                with pr3: proj["link"] = st.text_input("Link",         value=proj.get("link",""), key=f"p_link_{i}", placeholder="github.com/you/project")
                bullets = proj.get("bullets", ["",""])
                while len(bullets) < 2: bullets.append("")
                new_pb = []
                for j, b in enumerate(bullets):
                    new_pb.append(st.text_input(f"Bullet {j+1}", value=b, key=f"p_bullet_{i}_{j}",
                                                placeholder="What it does + tech + impact", label_visibility="collapsed"))
                proj["bullets"] = new_pb
                pc1, pc2 = st.columns([3,1])
                with pc1:
                    if st.button("AI Write Bullets", key=f"ai_proj_{i}", use_container_width=True):
                        with st.spinner("Writing…"):
                            text, err = ai_suggest_project_bullets(proj.get("name","Project"), proj.get("tech",""), st.session_state["rb_jd_text"])
                        if err: st.error(err)
                        else:
                            parsed = [ln.strip().lstrip("-").strip() for ln in text.splitlines() if ln.strip()][:3]
                            while len(parsed) < 2: parsed.append("")
                            proj["bullets"] = parsed[:3]
                            st.session_state["_pending_rescore"] = True
                            st.rerun()
                with pc2:
                    if st.button("Remove", key=f"del_proj_{i}", use_container_width=True):
                        st.session_state["rb_projects"].pop(i)
                        st.rerun()

        section_header("Education")
        if st.button("+ Add Education", key="add_edu"):
            st.session_state["rb_education"].append({"degree":"","institution":"","year":"","gpa":""})
            st.rerun()

        for i, edu in enumerate(st.session_state["rb_education"]):
            lbl = f"{edu['degree']} — {edu['institution']}" if edu["degree"] else f"Education {i+1}"
            with st.expander(lbl, expanded=True):
                ed1, ed2, ed3, ed4, ed5 = st.columns([2,2,1,1,1])
                with ed1: edu["degree"]      = st.text_input("Degree",      value=edu.get("degree",""),      key=f"ed_deg_{i}",  placeholder="B.Tech Computer Science")
                with ed2: edu["institution"] = st.text_input("Institution", value=edu.get("institution",""), key=f"ed_inst_{i}", placeholder="IIT Bombay")
                with ed3: edu["year"]        = st.text_input("Year",        value=edu.get("year",""),        key=f"ed_yr_{i}",   placeholder="2021–2025")
                with ed4: edu["gpa"]         = st.text_input("GPA",         value=edu.get("gpa",""),         key=f"ed_gpa_{i}",  placeholder="8.9 / 10")
                with ed5:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("Remove", key=f"del_edu_{i}", use_container_width=True):
                        st.session_state["rb_education"].pop(i)
                        st.rerun()

        section_header("Certifications")
        if st.button("+ Add Certification", key="add_cert"):
            st.session_state["rb_certs"].append({"name":"","issuer":"","year":""})
            st.rerun()

        for i, cert in enumerate(st.session_state["rb_certs"]):
            cr1, cr2, cr3, cr4 = st.columns([3,2,1,1])
            with cr1: cert["name"]   = st.text_input("Certificate", value=cert.get("name",""),   key=f"cr_name_{i}", placeholder="AWS Solutions Architect")
            with cr2: cert["issuer"] = st.text_input("Issuer",      value=cert.get("issuer",""), key=f"cr_iss_{i}",  placeholder="Amazon")
            with cr3: cert["year"]   = st.text_input("Year",        value=cert.get("year",""),   key=f"cr_yr_{i}",   placeholder="2024")
            with cr4:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("Remove", key=f"del_cert_{i}", use_container_width=True):
                    st.session_state["rb_certs"].pop(i)
                    st.rerun()

        if _ATS_AVAILABLE:
            st.markdown("---")
            section_header("ATS Score")
            ac1, ac2 = st.columns([1,2])
            with ac1:
                if st.button("Score My Resume", key="score_build_btn", use_container_width=True, type="primary"):
                    _run_ats_score()
                    st.rerun()
            with ac2:
                sn = st.session_state.get("rb_ats_result", {})
                if sn:
                    sc_val = sn["total"]; col = sn["color"]; lbl = sn["label"]
                    st.markdown(
                        f"<div style='background:{col}12;border:1px solid {col}44;border-radius:10px;"
                        f"padding:10px 16px;display:flex;align-items:center;gap:12px;'>"
                        f"<span style='font-size:2rem;font-weight:900;color:{col};'>{sc_val}</span>"
                        f"<div><div style='font-size:0.75rem;color:{col};font-weight:800;'>{lbl}</div>"
                        f"<div style='font-size:0.68rem;color:#7a9ec0;'>Full breakdown above</div></div></div>",
                        unsafe_allow_html=True,
                    )
                else:
                    st.caption("Click the button to score your current resume")

    # ══════════════════════════════════════════════════════════════════════
    # TAB 2 — AI ANALYSE
    # ══════════════════════════════════════════════════════════════════════
    with tab_ai:
        ai_left, ai_right = st.columns([1,1], gap="large")

        with ai_left:
            st.markdown("#### Target Internship JD")
            jd_r2, jd_c2 = st.columns(2)
            with jd_r2: t2_role = st.text_input("Role",    value=st.session_state.get("rb_jd_role",""),    key="t2_jd_role",    placeholder="Software Engineering Intern")
            with jd_c2: t2_co   = st.text_input("Company", value=st.session_state.get("rb_jd_company",""), key="t2_jd_company", placeholder="Google India")
            t2_jd = st.text_area("Job Description", value=st.session_state.get("rb_jd_text",""),
                                  height=260, key="t2_jd_text", placeholder="Paste full JD here…")
            if st.button("Run Analysis", key="run_analysis", use_container_width=True, type="primary"):
                resume_text = build_resume_text()
                if len(resume_text.strip()) < 50:
                    st.error("Resume is empty — fill in the Build tab first.")
                elif not t2_jd.strip():
                    st.error("Please paste the job description above.")
                else:
                    with st.spinner("Analysing resume vs JD…"):
                        analysis, err = ai_full_analysis(resume_text, t2_role, t2_jd)
                    if err: st.error(f"Analysis error: {err}")
                    else:
                        st.session_state["rb_analysis"] = analysis
                        st.success("Analysis complete — see results on the right.")

        with ai_right:
            analysis = st.session_state.get("rb_analysis", {})
            if not analysis:
                st.markdown(
                    "<div style='background:rgba(0,0,0,0.2);border:2px dashed rgba(74,144,217,0.15);"
                    "border-radius:12px;padding:52px 24px;text-align:center;'>"
                    "<div style='color:#3a5575;font-size:0.9rem;font-weight:700;'>Analysis results appear here</div>"
                    "<div style='color:#2a3a50;font-size:0.77rem;margin-top:6px;'>Paste JD and click Run Analysis</div></div>",
                    unsafe_allow_html=True,
                )
            else:
                ms = analysis.get("match_score", 0)
                at = analysis.get("ats_score", 0)
                st.markdown(
                    f"<div style='display:grid;grid-template-columns:1fr 1fr;gap:10px;margin-bottom:14px;'>"
                    f"<div style='background:rgba(0,0,0,0.35);border:1px solid {score_color(ms)}44;border-radius:12px;padding:14px;text-align:center;'>"
                    f"<div style='font-size:2.6rem;font-weight:900;color:{score_color(ms)};line-height:1;'>{ms}</div>"
                    f"<div style='font-size:0.6rem;color:{score_color(ms)};font-weight:800;text-transform:uppercase;margin:4px 0 2px;'>{score_label(ms)}</div>"
                    f"<div style='font-size:0.72rem;color:#7a9ec0;'>JD Match</div></div>"
                    f"<div style='background:rgba(0,0,0,0.35);border:1px solid {score_color(at)}44;border-radius:12px;padding:14px;text-align:center;'>"
                    f"<div style='font-size:2.6rem;font-weight:900;color:{score_color(at)};line-height:1;'>{at}</div>"
                    f"<div style='font-size:0.6rem;color:{score_color(at)};font-weight:800;text-transform:uppercase;margin:4px 0 2px;'>{score_label(at)}</div>"
                    f"<div style='font-size:0.72rem;color:#7a9ec0;'>ATS Score</div></div></div>",
                    unsafe_allow_html=True,
                )
                sec = analysis.get("section_scores", {})
                if sec:
                    bars = "<div style='margin-bottom:14px;'><div style='font-size:0.67rem;color:#7ec8f7;font-weight:700;text-transform:uppercase;letter-spacing:0.5px;margin-bottom:8px;'>Section Scores</div>"
                    for sname, sc_val in sec.items():
                        pct = int(sc_val) * 10
                        col = score_color(pct)
                        bars += (f"<div style='display:flex;align-items:center;gap:8px;margin-bottom:6px;'>"
                                 f"<div style='width:80px;font-size:0.7rem;color:#9ab8d8;text-transform:capitalize;'>{sname}</div>"
                                 f"<div style='flex:1;background:rgba(255,255,255,0.05);border-radius:4px;height:8px;overflow:hidden;'>"
                                 f"<div style='width:{pct}%;height:100%;background:{col};border-radius:4px;'></div></div>"
                                 f"<div style='width:30px;font-size:0.68rem;color:{col};font-weight:700;text-align:right;'>{sc_val}/10</div></div>")
                    bars += "</div>"
                    st.markdown(bars, unsafe_allow_html=True)

                missing_kw = analysis.get("missing_keywords", [])
                matched_kw = analysis.get("matched_keywords", [])
                if matched_kw or missing_kw:
                    pills = "<div style='margin-bottom:12px;'>"
                    if matched_kw:
                        pills += "<div style='font-size:0.67rem;color:#4ade80;font-weight:700;margin-bottom:5px;'>In your resume</div><div style='margin-bottom:8px;'>"
                        for k in matched_kw[:10]: pills += pill_badge(k, "#4ade80")
                        pills += "</div>"
                    if missing_kw:
                        pills += "<div style='font-size:0.67rem;color:#f87171;font-weight:700;margin-bottom:5px;'>Missing keywords</div><div style='margin-bottom:8px;'>"
                        for k in missing_kw[:14]: pills += pill_badge(k, "#f87171")
                        pills += "</div>"
                    pills += "</div>"
                    st.markdown(pills, unsafe_allow_html=True)

                t_s, t_g, t_q = st.tabs(["Strengths", "Gaps", "Quick Wins"])
                with t_s:
                    for s in analysis.get("strengths", ["—"]): st.markdown(f"- {s}")
                with t_g:
                    for g in analysis.get("critical_gaps", ["—"]): st.markdown(f"- {g}")
                with t_q:
                    for q in analysis.get("quick_wins", ["—"]): st.markdown(f"- {q}")

                with st.expander("LinkedIn Headline + Interview Questions", expanded=False):
                    ll = analysis.get("linkedin_headline","")
                    if ll:
                        st.markdown("**Suggested LinkedIn Headline:**")
                        st.code(ll, language="")
                    iqs = analysis.get("interview_questions", [])
                    if iqs:
                        st.markdown("**Prepare for these questions:**")
                        for idx, q in enumerate(iqs, 1): st.markdown(f"{idx}. {q}")

    # ══════════════════════════════════════════════════════════════════════
    # TAB 3 — PREVIEW & EXPORT
    # ══════════════════════════════════════════════════════════════════════
    with tab_preview:
        prev_left, prev_right = st.columns([1,1], gap="large")

        with prev_left:
            st.markdown("#### Template")
            template_choice = st.radio(
                "Choose template", ["Modern","Professional","Minimal"],
                index=["Modern","Professional","Minimal"].index(st.session_state["rb_template"]),
                horizontal=True, key="tpl_radio", label_visibility="collapsed",
            )
            st.session_state["rb_template"] = template_choice

            tpl_info = {
                "Modern":       ("Dark header, blue accents — great for tech roles", "#1e293b"),
                "Professional": ("Classic serif, navy accents — ideal for consulting/finance", "#1a4e8a"),
                "Minimal":      ("Clean grey layout — universally ATS safe", "#374151"),
            }
            info, color = tpl_info[template_choice]
            st.markdown(
                f"<div style='background:rgba(0,0,0,0.25);border:1px solid {color}44;"
                f"border-radius:8px;padding:10px 14px;font-size:0.78rem;color:#9ab8d8;'>"
                f"<b style='color:#e8eef8;'>{template_choice}</b> — {info}</div>",
                unsafe_allow_html=True,
            )

            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown("#### Download")
            resume_text  = build_resume_text()
            html_content = render_html_template(template_choice)
            cslug = (st.session_state["rb_jd_company"] or "resume").lower().replace(" ","_")
            rslug = (st.session_state["rb_name"]        or "me").lower().replace(" ","_")
            base  = f"resume_{rslug}_{cslug}"

            d1, d2 = st.columns(2)
            with d1:
                st.download_button("Download HTML", data=html_content.encode("utf-8"),
                                   file_name=f"{base}.html", mime="text/html",
                                   use_container_width=True, key="dl_html")
            with d2:
                st.download_button("Download TXT", data=resume_text.encode("utf-8"),
                                   file_name=f"{base}.txt", mime="text/plain",
                                   use_container_width=True, key="dl_txt")

            st.markdown("<br>**Inject into your DOCX:**")
            st.caption("Upload your original DOCX — text will be rewritten, formatting preserved.")
            inject_file = st.file_uploader("Upload DOCX", type=["docx"],
                                           key="dl_docx_inject", label_visibility="collapsed")
            if inject_file:
                updated = inject_text_into_docx(inject_file.getvalue(), resume_text)
                st.download_button("Download Updated DOCX", data=updated, file_name=f"{base}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   use_container_width=True, key="dl_docx")

            st.markdown(
                "<div style='background:rgba(74,144,217,0.07);border:1px solid rgba(74,144,217,0.18);"
                "border-radius:10px;padding:12px 16px;margin-top:14px;font-size:0.78rem;color:#9ab8d8;line-height:1.9;'>"
                "<b style='color:#7ec8f7;'>To save as PDF:</b><br>"
                "1. Download HTML &nbsp; 2. Open in Chrome &nbsp; 3. Ctrl+P → Save as PDF"
                "</div>",
                unsafe_allow_html=True,
            )

        with prev_right:
            st.markdown("#### Live Preview")
            st.caption(f"Template: {template_choice} — updates as you edit")
            st.components.v1.html(render_html_template(template_choice), height=820, scrolling=True)
        
    # ══════════════════════════════════════════════════════════════════════
    # TAB 4 — SKILL ROADMAP
    # ══════════════════════════════════════════════════════════════════════
    with tab_roadmap:
        st.markdown("#### Personalised Skill Roadmap")

        rm_role = st.text_input(
            "Target role",
            value=st.session_state.get("rb_jd_role", ""),
            placeholder="ML Engineer at Amazon India",
            key="rm_role",
        )

        rm_weeks = st.slider("Timeline (weeks)", 4, 16, 8, key="rm_weeks")

        if st.button("Generate Roadmap", type="primary", key="gen_roadmap_btn"):
            current_skills = st.session_state.get("rb_skills", "")
            jd_text        = st.session_state.get("rb_jd_text", "")

            if not rm_role.strip():
                st.error("Enter target role")
            else:
                with st.spinner("Building your roadmap…"):
                    result, err = ai_skill_roadmap(
                        rm_role.strip(), rm_weeks, current_skills, jd_text
                    )
                if err:
                    st.error(err)
                elif not result:
                    st.error("Could not generate roadmap.")
                else:
                    st.session_state["roadmap_result"] = result

        if st.session_state.get("roadmap_result"):
            st.markdown(st.session_state["roadmap_result"])
            st.download_button(
                "Download Roadmap",
                st.session_state["roadmap_result"].encode("utf-8"),
                "skill_roadmap.md",
                "text/markdown",
                key="dl_roadmap",
            )
        else:
            st.caption("Generate a roadmap based on your target role, current skills, and JD.")


if __name__ == "__main__":
    main()